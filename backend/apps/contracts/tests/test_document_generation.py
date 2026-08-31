"""Tests for invoice document generation (Slice 1 — Invoice RU/EN).

Three layers:
  * Pure context-builder unit tests (no DB; SimpleNamespace mocks).
  * Render smoke tests (fills the shipped templates; asserts no leftover tags).
  * API integration tests (real DB) for the /sales/{id}/document/ endpoint.

Reuses the fixture helpers from test_contract_sale_api.
"""
import re
import zipfile
from datetime import date
from decimal import Decimal
from io import BytesIO
from types import SimpleNamespace
from unittest import mock

import openpyxl
from docx import Document
from docx.shared import Cm, Emu, Mm
from rest_framework.test import APIClient

from django.core.exceptions import ValidationError
from django.test import SimpleTestCase, TestCase

from apps.core.models import Country, ShipmentStatusType
from apps.export.models import Shipment, ShipmentFirmSplit
from apps.contracts.document_templates import registry as tpl_registry
from apps.contracts.document_templates.registry import get_spec
from apps.contracts.models import DocumentLayoutSetting
from apps.contracts.services import document_context as ctx
from apps.contracts.services import document_highlight as highlight
from apps.contracts.services import document_render as render
from apps.contracts.tests.test_contract_sale_api import (
    _SeededPermsMixin,
    _make_contract,
    _make_export_firm,
    _make_import_firm,
    _make_invoice,
    _make_season,
    _make_user,
)


def _make_country(name='Kazakhstan', code='KZ'):
    country, _ = Country.objects.get_or_create(
        code=code, defaults={'name_tk': name, 'name_ru': name, 'name_en': name},
    )
    return country


def _make_packed_shipment(season, import_firm, status_code='draft', code='0101777/25',
                          driver_name='Ahmet A.', truck_plate='BR1427LB', with_country=True):
    """A shipment with complete packing + the destination fields the Documents page
    gates on (import firm, country, driver, plate). Pass blanks / with_country=False
    to simulate an incomplete truck."""
    status, _ = ShipmentStatusType.objects.get_or_create(
        code=status_code,
        defaults={'name_tk': status_code, 'name_en': status_code.title(),
                  'name_ru': status_code, 'step_order': 0, 'phase': 'PREP'},
    )
    return Shipment.objects.create(
        shipment_code=code, date='2025-10-01', season=season, status=status,
        import_firm=import_firm, weight_gross=Decimal('10720'), weight_net=Decimal('9000'),
        box_count=1800, pallet_count=16,
        driver_name=driver_name, truck_plate=truck_plate,
        country=_make_country() if with_country else None,
    )


def _preset(**kw):
    """A PackingTemplate stand-in (whole-truck values), for the CMR."""
    return SimpleNamespace(
        net_kg=kw.get('net_kg'), gross_kg=kw.get('gross_kg'),
        box_count=kw.get('box_count'), pallet_count=kw.get('pallet_count'),
        pallet_weight_kg=kw.get('pallet_weight_kg'),
    )


def _mock_invoice(*, with_shipment=True, truck_template=None,
                  quantity_kg=Decimal('9000'), packing=None):
    """Build a SimpleNamespace invoice mirroring the real ORM attributes used.

    ``truck_template`` = whole-truck PackingTemplate on the shipment (drives the CMR).
    ``packing`` = the firm's explicit packing on the sale (gross_kg/box_count/…), the
    values copied from the template share and printed on the Invoice.
    """
    seller = SimpleNamespace(
        name_ru='Х.О «Датлы миве»', name_en='Datly miwe LLC', name_tk='',
        address_ru='г. Ашгабат', address_en='Ashgabat', address_tk='',
        bank_details_ru='Банк: АКБТ', bank_details_en='Bank: SCBT', bank_details_tk='',
    )
    country = SimpleNamespace(name_ru='Узбекистан', name_en='Uzbekistan', name_tk='Özbegistan')
    buyer = SimpleNamespace(
        name_company='ООО TRUST', address='г. Ташкент', bank_details='ИНН: 311270964',
        country=country,
    )
    shipment = SimpleNamespace(
        weight_net=Decimal('9000'), weight_gross=Decimal('10720'), box_count=1800,
        pallet_count=16, packaging_kg=Decimal('300'), pallet_weight_kg=Decimal('300'),
        truck_plate='BR1427LB', trailer_id=5311, driver_name='Ahmet A.', country=country,
        packing_template=truck_template,
    ) if with_shipment else None
    contract = SimpleNamespace(
        contract_number='93/26-DM-EXP', start_date=date(2026, 3, 16),
        export_firm=seller, import_firm=buyer, incoterm='FCA',
    )
    pk = packing or {}
    return SimpleNamespace(
        invoice_number=118, invoice_date=date(2026, 3, 16), contract=contract,
        shipment=shipment, export_firm=seller, import_firm=buyer, incoterm='FCA',
        quantity_kg=quantity_kg, price_per_kg=Decimal('0.87'),
        total_usd=Decimal('7830'),
        gross_kg=pk.get('gross_kg'), box_count=pk.get('box_count'),
        pallet_count=pk.get('pallet_count'), pallet_weight_kg=pk.get('pallet_weight_kg'),
    )


def _line(name, qty, price, *, gross=None, boxes=None, hs=''):
    """A ContractSaleLineItem stand-in for the invoice builder."""
    q, p = Decimal(qty), Decimal(price)
    return SimpleNamespace(
        product_name=name, hs_code=hs, quantity_kg=q, price_per_kg=p,
        gross_kg=Decimal(gross) if gross else None, box_count=boxes, total_usd=q * p,
    )


def _invoice_with_lines(*lines):
    """A mock invoice carrying explicit line items (``.line_items.all()``)."""
    inv = _mock_invoice()
    inv.line_items = SimpleNamespace(all=lambda: list(lines))
    return inv


def _mock_firm(name_ru, name_en, address_ru, address_en):
    return SimpleNamespace(
        name_ru=name_ru, name_en=name_en, name_tk='',
        address_ru=address_ru, address_en=address_en, address_tk='',
        bank_details_ru='', bank_details_en='', bank_details_tk='',
    )


def _mock_shipment(*, firms=None, truck_template=None):
    """A SimpleNamespace shipment mirroring the ORM attributes the CMR builder reads.

    ``firms`` = the export firms on the truck (default one); each becomes a firm
    split and a matching invoice. ``.firm_splits`` / ``.sales`` expose ``.all()``.
    """
    if firms is None:
        firms = [_mock_firm('Х.О «Датлы миве»', 'Datly miwe LLC', 'г. Ашгабат', 'Ashgabat')]
    country = SimpleNamespace(name_ru='Узбекистан', name_en='Uzbekistan', name_tk='Özbegistan')
    buyer = SimpleNamespace(
        name_company='ООО TRUST', address='г. Ташкент', bank_details='ИНН: 311270964', country=country,
    )
    splits = [SimpleNamespace(export_firm=firm) for firm in firms]
    sales = [SimpleNamespace(invoice_number=118 + i, invoice_date=date(2026, 3, 16))
             for i in range(len(firms))]
    return SimpleNamespace(
        shipment_code='0316118/25',
        firm_splits=SimpleNamespace(all=lambda: splits),
        sales=SimpleNamespace(all=lambda: sales),
        import_firm=buyer, packing_template=truck_template,
        weight_net=Decimal('9000'), weight_gross=Decimal('10720'), box_count=1800,
        pallet_count=16, packaging_kg=Decimal('300'), pallet_weight_kg=Decimal('300'),
        truck_plate='BR1427LB', trailer_id=5311, driver_name='Ahmet A.', date=date(2026, 3, 16),
    )


class InvoiceContextBuilderTest(SimpleTestCase):
    """Pure builder: formatting, language selection, shipment fallback."""

    def test_ru_context_values_and_formatting(self):
        c = ctx.build_invoice_context(_mock_invoice(), 'ru')
        self.assertEqual(c['invoice_no'], '118')
        self.assertEqual(c['invoice_date'], '16.03.2026')
        self.assertEqual(c['contract_line'], '93/26-DM-EXP, 16.03.2026')
        self.assertEqual(c['seller_name'], 'Х.О «Датлы миве»')
        self.assertEqual(c['buyer_name'], 'ООО TRUST')
        # RU convention: space-thousands + comma-decimal
        self.assertEqual(c['total_sum'], '7 830,00')
        item = c['line_items'][0]
        self.assertEqual(item['name'], 'Помидор свежий')
        self.assertEqual(item['code'], ctx.TOMATO_HS_CODE)
        self.assertEqual(item['gross'], '10 720')
        self.assertEqual(item['net'], '9 000')
        self.assertEqual(item['price'], '0,87')
        self.assertIn('2026', c['country_origin'])

    def test_en_keeps_english_number_format_and_firm_columns(self):
        c = ctx.build_invoice_context(_mock_invoice(), 'en')
        self.assertEqual(c['seller_name'], 'Datly miwe LLC')
        self.assertEqual(c['line_items'][0]['name'], 'Fresh tomatoes')
        self.assertIn('harvest', c['country_origin'])
        # EN convention unchanged: comma-thousands + dot-decimal
        self.assertEqual(c['total_sum'], '7,830.00')
        self.assertEqual(c['line_items'][0]['gross'], '10,720')
        self.assertEqual(c['line_items'][0]['price'], '0.87')

    def test_place_loading_override(self):
        # blank by default, filled from generate-time overrides
        self.assertEqual(ctx.build_invoice_context(_mock_invoice(), 'ru')['place_loading'], '')
        c = ctx.build_invoice_context(_mock_invoice(), 'ru', {'place_loading': 'Kaka'})
        self.assertEqual(c['place_loading'], 'Kaka')

    def test_multi_line_from_line_items(self):
        # explicit line items → multiple rows; total_sum = their sum (not the sale's)
        inv = _invoice_with_lines(
            _line('Сорт А', '5000', '1.00'),
            _line('Сорт Б', '4000', '0.80'),
        )
        c = ctx.build_invoice_context(inv, 'ru')
        self.assertEqual(len(c['line_items']), 2)
        self.assertEqual([li['name'] for li in c['line_items']], ['Сорт А', 'Сорт Б'])
        self.assertEqual([li['n'] for li in c['line_items']], ['1', '2'])
        self.assertEqual(c['line_items'][0]['price'], '1')      # 1.00 → trimmed
        self.assertEqual(c['total_sum'], '8 200,00')            # 5000·1 + 4000·0.8

    def test_line_name_falls_back_to_default(self):
        inv = _invoice_with_lines(_line('', '9000', '0.87'))
        c = ctx.build_invoice_context(inv, 'ru')
        self.assertEqual(c['line_items'][0]['name'], 'Помидор свежий')
        self.assertEqual(c['line_items'][0]['code'], ctx.TOMATO_HS_CODE)

    def test_no_line_items_uses_single_synthesized_line(self):
        # a sale with an empty line_items manager → the classic single line
        inv = _invoice_with_lines()  # .all() == []
        c = ctx.build_invoice_context(inv, 'ru')
        self.assertEqual(len(c['line_items']), 1)
        self.assertEqual(c['line_items'][0]['name'], 'Помидор свежий')
        self.assertEqual(c['total_sum'], '7 830,00')           # the sale's total_usd

    def test_multi_line_gross_boxes_allocated_by_weight(self):
        # lines carry no gross/box → split the sale's whole-truck gross(10720) /
        # boxes(1800, from the shipment) by each line's weight share (no blank cell).
        inv = _invoice_with_lines(_line('A', '6000', '1.00'), _line('B', '3000', '1.00'))
        c = ctx.build_invoice_context(inv, 'ru')
        self.assertEqual(c['line_items'][0]['pieces'], '1200')   # 1800 · 6000/9000
        self.assertEqual(c['line_items'][1]['pieces'], '600')    # 1800 · 3000/9000
        self.assertEqual(c['line_items'][0]['gross'], '7 147')   # 10720 · 2/3, rounded
        self.assertEqual(c['line_items'][1]['gross'], '3 573')   # 10720 · 1/3

    def test_falls_back_to_invoice_qty_when_no_shipment(self):
        c = ctx.build_invoice_context(_mock_invoice(with_shipment=False), 'ru')
        # net falls back to invoice.quantity_kg; gross/transport empty
        self.assertEqual(c['line_items'][0]['net'], '9 000')
        self.assertEqual(c['line_items'][0]['gross'], '')
        self.assertEqual(c['transport'], '')
        self.assertEqual(c['pallet_note'], '')

    def test_null_invoice_number_renders_blank_not_none(self):
        # invoice_number is nullable (bridge sale / not yet numbered) — the
        # document must show blank, never the literal "None".
        inv = _mock_invoice()
        inv.invoice_number = None
        c = ctx.build_invoice_context(inv, 'ru')
        self.assertEqual(c['invoice_no'], '')
        self.assertEqual(ctx.invoice_filename_fields(inv)['invoice_number'], 'NA')

    def test_per_firm_packing_prints_explicit_share(self):
        """The invoice prints the firm's explicit packing (copied from the share)."""
        c = ctx.build_invoice_context(
            _mock_invoice(quantity_kg=Decimal('10000'),
                          packing={'gross_kg': Decimal('11373'), 'box_count': 1618,
                                   'pallet_count': Decimal('18'), 'pallet_weight_kg': Decimal('229')}),
            'ru',
        )
        item = c['line_items'][0]
        self.assertEqual(item['net'], '10 000')    # firm's own weight (official)
        self.assertEqual(item['gross'], '11 373')  # the firm's share gross, not the truck
        self.assertEqual(item['pieces'], '1618')

    def test_per_firm_packing_falls_back_to_shipment(self):
        """With no per-firm packing set, gross/boxes fall back to the shipment."""
        c = ctx.build_invoice_context(_mock_invoice(), 'ru')  # no packing dict
        item = c['line_items'][0]
        self.assertEqual(item['gross'], '10 720')  # shipment.weight_gross
        self.assertEqual(item['pieces'], '1800')   # shipment.box_count


class PackingGuardTest(SimpleTestCase):
    """missing_packing_on: raw cells OR an applied PackingTemplate satisfy the guard."""

    def test_complete_raw_cells_pass(self):
        self.assertEqual(ctx.missing_packing_on(_mock_shipment()), [])

    def test_no_shipment_all_missing(self):
        self.assertEqual(len(ctx.missing_packing_on(None)), 4)

    def test_missing_raw_cell_reported(self):
        ship = _mock_shipment()
        ship.box_count = None
        self.assertEqual(ctx.missing_packing_on(ship), ['box_count'])

    def test_template_fills_null_raw_cells(self):
        # Template-configured truck: raw cells null, but the template supplies all.
        ship = _mock_shipment(truck_template=_preset(
            net_kg=Decimal('18000'), gross_kg=Decimal('20450'),
            box_count=2984, pallet_count=Decimal('33'), pallet_weight_kg=Decimal('446'),
        ))
        ship.weight_gross = ship.weight_net = ship.box_count = ship.pallet_count = None
        self.assertEqual(ctx.missing_packing_on(ship), [])

    def test_neither_source_blocks(self):
        ship = _mock_shipment()  # no template
        ship.weight_gross = ship.weight_net = ship.box_count = ship.pallet_count = None
        self.assertEqual(sorted(ctx.missing_packing_on(ship)), sorted(ctx.REQUIRED_PACKING_FIELDS))


class CmrContextBuilderTest(SimpleTestCase):
    """Pure CMR builder (truck-level): gross-with-pallet math, transport, refs."""

    def test_ru_cmr_values(self):
        c = ctx.build_cmr_context(_mock_shipment(), 'ru')
        self.assertEqual(c['sender_name'], 'Х.О «Датлы миве»')
        # forwarder is the export firm(s) (same as the sender box)
        self.assertEqual(c['forwarder'], 'Х.О «Датлы миве»')
        self.assertEqual(c['consignee_name'], 'ООО TRUST')
        self.assertEqual(c['country_dispatch'], 'Туркменистан')
        self.assertEqual(c['cargo_name'], 'Помидоры свежие')
        self.assertEqual(c['boxes'], '1800')
        self.assertEqual(c['pallets'], '16')
        # weight_gross is BRUT (with pallet): with = 10720, without = 10720 - 300
        self.assertEqual(c['gross_with_pallet'], '10 720')
        self.assertEqual(c['gross_without_pallet'], '10 420')
        self.assertEqual(c['net'], '9 000')
        self.assertIn('118', c['invoice_refs'])
        self.assertIn('Ahmet A.', c['transport'])
        # generate-time fields stay blank when no overrides are supplied
        self.assertEqual(c['tir_carnet'], '')
        self.assertEqual(c['place_loading'], '')

    def test_multi_firm_lists_all_senders(self):
        firms = [
            _mock_firm('Х.О «Датлы миве»', 'Datly miwe LLC', 'г. Ашгабат', 'Ashgabat'),
            _mock_firm('Х.О «Ýigit»', 'Yigit LLC', 'г. Мары', 'Mary'),
        ]
        c = ctx.build_cmr_context(_mock_shipment(firms=firms), 'ru')
        # both export firms appear in the single sender box (newline-joined)
        self.assertIn('Датлы миве', c['sender_name'])
        self.assertIn('Ýigit', c['sender_name'])
        # both invoices referenced on the one truck CMR
        self.assertIn('118', c['invoice_refs'])
        self.assertIn('119', c['invoice_refs'])

    def test_generate_time_overrides(self):
        c = ctx.build_cmr_context(
            _mock_shipment(), 'ru',
            {'place_loading': 'Dusak', 'tir_carnet': 'RU 82345678'},
        )
        self.assertEqual(c['place_loading'], 'Dusak')
        self.assertEqual(c['tir_carnet'], 'RU 82345678')

    def test_en_cmr_localization(self):
        c = ctx.build_cmr_context(_mock_shipment(), 'en')
        self.assertEqual(c['cargo_name'], 'FRESH TOMATOES')
        self.assertEqual(c['country_dispatch'], 'Turkmenistan')
        self.assertEqual(c['gross_with_pallet'], '10,720')

    def test_null_invoice_number_skipped_in_refs(self):
        # a bridge sale with no invoice_number must not print "None" on the CMR
        ship = _mock_shipment()
        ship.sales = SimpleNamespace(all=lambda: [
            SimpleNamespace(invoice_number=118, invoice_date=date(2026, 3, 16)),
            SimpleNamespace(invoice_number=None, invoice_date=date(2026, 3, 16)),
        ])
        c = ctx.build_cmr_context(ship, 'ru')
        self.assertIn('118', c['invoice_refs'])
        self.assertNotIn('None', c['invoice_refs'])


class CmrPresetTest(SimpleTestCase):
    """CMR reads the whole-truck template on the shipment; BRUT = gross WITH pallet."""

    def test_whole_truck_template_drives_cmr(self):
        # gross-net row 152 right block: BRUT 20450, NET 18000, boxes 2984, 33 pal, pallet 446.
        truck = _preset(
            net_kg=Decimal('18000'), gross_kg=Decimal('20450'), box_count=2984,
            pallet_count=Decimal('33'), pallet_weight_kg=Decimal('446'),
        )
        c = ctx.build_cmr_context(_mock_shipment(truck_template=truck), 'ru')
        self.assertEqual(c['boxes'], '2984')
        self.assertEqual(c['pallets'], '33')                 # _num drops the .0
        self.assertEqual(c['net'], '18 000')
        self.assertEqual(c['pallet_weight'], '446')
        self.assertEqual(c['gross_with_pallet'], '20 450')   # BRUT as-is
        self.assertEqual(c['gross_without_pallet'], '20 004')  # BRUT − pallet weight

    def test_net_from_template_when_gross_absent(self):
        # net_kg resolves independently of gross_kg: template supplies net, raw
        # cell is null → net still fills (mirrors the packing guard). Regression
        # for the guard/builder divergence caught in review.
        ship = _mock_shipment(truck_template=_preset(net_kg=Decimal('18000')))
        ship.weight_net = None
        c = ctx.build_cmr_context(ship, 'ru')
        self.assertEqual(c['net'], '18 000')


class LetterContextBuilderTest(SimpleTestCase):
    """Pure builders for the CT-1 / phyto / customs request letters."""

    def test_ct1_needs_only_firm_and_contract(self):
        c = ctx.build_ct1_context(_mock_invoice(), 'ru')
        self.assertEqual(c['firm_name'], 'Х.О «Датлы миве»')
        self.assertIn('93/26-DM-EXP', c['contract_line'])
        self.assertEqual(c['product'], 'Свежие Помидоры')

    def test_fito_resolves_country_weight_boxes(self):
        c = ctx.build_fito_context(_mock_invoice(), 'ru')
        self.assertEqual(c['country'], 'Узбекистан')
        self.assertEqual(c['net'], '9 000')   # RU space-thousands
        self.assertEqual(c['boxes'], '1800')

    def test_customs_is_turkmen_with_both_firms(self):
        c = ctx.build_customs_context(_mock_invoice(), 'tk')
        self.assertEqual(c['buyer_name'], 'ООО TRUST')
        self.assertEqual(c['country'], 'Özbegistan')  # tk country name
        self.assertIn('93/26-DM-EXP', c['contract_line'])

    def test_country_falls_back_when_no_shipment(self):
        # No shipment → resolver falls back to buyer firm's country
        c = ctx.build_fito_context(_mock_invoice(with_shipment=False), 'ru')
        self.assertEqual(c['country'], 'Узбекистан')


class InvoiceRenderSmokeTest(TestCase):
    """Fill the shipped templates and assert clean, value-bearing output.

    ``TestCase``, not ``SimpleTestCase``: rendering reads the document's saved
    ``DocumentLayoutSetting`` row. The mocks themselves still need no fixtures.
    """

    def _text(self, data: bytes) -> str:
        doc = Document(BytesIO(data))
        parts = [p.text for p in doc.paragraphs]
        for t in doc.tables:
            for row in t.rows:
                parts.append(' | '.join(c.text for c in row.cells))
        return '\n'.join(parts)

    def test_render_docx_ru_and_en(self):
        expected_total = {'invoice_ru': '7 830,00', 'invoice_en': '7,830.00'}
        for key in ('invoice_ru', 'invoice_en'):
            data, filename, content_type = render.generate(key, _mock_invoice(), 'docx')
            text = self._text(data)
            self.assertNotIn('{{', text, f'{key}: unrendered tag')
            self.assertNotIn('{%', text, f'{key}: unrendered tag')
            self.assertIn('118', text)
            self.assertIn(expected_total[key], text)
            self.assertTrue(filename.endswith('.docx'))
            self.assertEqual(content_type, render.DOCX_CONTENT_TYPE)
            self.assertIn('118', filename)

    def test_render_multi_line_invoice(self):
        # the docxtpl row loop emits one table row per line item; both render.
        inv = _invoice_with_lines(
            _line('Pomidor A', '5000', '1.00'),
            _line('Pomidor B', '4000', '0.80'),
        )
        data, _f, _ct = render.generate('invoice_ru', inv, 'docx')
        text = self._text(data)
        self.assertNotIn('{{', text, 'unrendered tag')
        self.assertNotIn('{%', text, 'unrendered loop tag')
        self.assertIn('Pomidor A', text)
        self.assertIn('Pomidor B', text)
        self.assertIn('8 200,00', text)   # summed total

    def _xlsx_text(self, data: bytes) -> str:
        wb = openpyxl.load_workbook(BytesIO(data))
        ws = wb.active
        return '\n'.join(
            str(c.value) for row in ws.iter_rows() for c in row if c.value is not None
        )

    def test_render_cmr_ru_and_en(self):
        # CMR is an xlsx print-overlay (not docx). Two firms on the truck → both
        # sender names must survive into the filled sheet (joined in the sender box).
        firms = [
            _mock_firm('Х.О «Датлы миве»', 'Datly miwe LLC', 'г. Ашгабат', 'Ashgabat'),
            _mock_firm('Х.О «Ýigit»', 'Yigit LLC', 'г. Мары', 'Mary'),
        ]
        for key, names in (('cmr_ru', ('Датлы миве', 'Ýigit')), ('cmr_en', ('Datly miwe', 'Yigit'))):
            data, filename, content_type = render.generate(key, _mock_shipment(firms=firms), 'docx')
            text = self._xlsx_text(data)
            self.assertIn('CMR_', filename)
            self.assertTrue(filename.endswith('.xlsx'), f'{key}: {filename}')
            for name in names:
                self.assertIn(name, text, f'{key}: sender {name!r} missing from render')
            self.assertEqual(content_type, render.XLSX_CONTENT_TYPE)

    def test_render_request_letters(self):
        # marker + the richer per-letter content added to match the Excel sheets:
        # ct1/fito carry the consignee block + weights; customs the truck table.
        expected = {
            'ct1_ru': ('СТ-1', ('Грузополучатель', 'ООО TRUST', '9 000')),
            'fito_ru': ('Фитосанитарный', ('Грузополучатель', 'BR1427LB', '9 000')),
            'customs_tk': ('ARZA', ('Ulag serişdeleriniň', 'BR1427LB', '10,720')),
        }
        for key, (marker, required) in expected.items():
            data, filename, content_type = render.generate(key, _mock_invoice(), 'docx')
            text = self._text(data)
            self.assertNotIn('{{', text, f'{key}: unrendered tag')
            self.assertNotIn('{%', text, f'{key}: unrendered tag')
            self.assertIn(marker, text)
            for token in required:
                self.assertIn(token, text, f'{key}: {token!r} missing from render')
            self.assertEqual(content_type, render.DOCX_CONTENT_TYPE)

    def test_unsupported_format_raises(self):
        with self.assertRaises(ValueError):
            render.generate('invoice_ru', _mock_invoice(), 'xlsx')

    def test_pdf_without_libreoffice_raises_clear_error(self):
        with mock.patch.object(render, '_libreoffice_bin', return_value=None):
            with self.assertRaises(render.DocumentRenderError):
                render.generate('invoice_ru', _mock_invoice(), 'pdf')


def _mock_seller():
    """ExportFirm stand-in: bilingual name/address/director + bank blob per language."""
    return SimpleNamespace(
        name_tk='Hemsaya H.J.', name_ru='Хемсая Х.Дж.',
        address_tk='Türkmenistan, Ahal w., Kaka etr.', address_ru='Туркменистан, Ахалская обл.',
        director='Директор Худайназаров Ы.',
        director_tk='Direktor Hudaýnazarow Ý.',
        bank_details_tk='Bank: Türkmenbaşy\nSWIFT: INVATM2X',
        bank_details_ru='Банк: Туркменбаши\nВал/счет: 23202\nSWIFT: INVATM2X',
    )


def _mock_contract(amount='7830.00', qty='9000', end=date(2026, 12, 31),
                   contact_person=None, contact_person_tk=None,
                   country_code='KZ', country_ru='Казахстан', country_tk='Gazagystan',
                   contract_date=None, price=None, start=date(2026, 3, 18)):
    """A SimpleNamespace Contract mirroring the ORM attributes the builder reads.

    Buyer is fidelity A: the flat single-value ImportFirm fields (name_company /
    address / bank_details blob). The director is the firm's ``contact_person``
    ("Director's Full Name", RU/Cyrillic) + ``contact_person_tk`` (TK/Latin), or a
    generate-time override when the RU form is blank.
    """
    country = SimpleNamespace(name_tk=country_tk, name_ru=country_ru, code=country_code)
    buyer = SimpleNamespace(
        name_company='TOO «Aranşy - KZ»', name_short='Aranşy',
        address='РК, Туркестанская обл., с. Первое Мая',
        bank_details='БИН 191040016779\nБИК HSBKKZKX\nр/с KZ97601A891001387241',
        contact_person=contact_person, contact_person_tk=contact_person_tk,
        country=country,
    )
    return SimpleNamespace(
        contract_number='108/26-YGT-EXP', start_date=start, end_date=end,
        contract_date=contract_date,
        planned_amount_usd=Decimal(amount) if amount is not None else None,
        planned_quantity_kg=Decimal(qty) if qty is not None else None,
        price_per_kg=Decimal(price) if price is not None else None,
        export_firm=_mock_seller(), import_firm=buyer,
    )


class ContractContextBuilderTest(SimpleTestCase):
    """Pure contract builder: financials, amount-in-words, dates, seller + buyer."""

    def test_destination_country_genitive_from_buyer_country(self):
        # §4.1/§4.2 name the buyer country's authorities in the genitive case —
        # taken from the code map, not from Country's nominative name_tk/name_ru.
        c = ctx.build_contract_context(_mock_contract(), 'ru')
        self.assertEqual(c['dest_country_gen_ru'], 'Казахстана')
        self.assertEqual(c['dest_country_gen_tk'], 'Gazagystanyň')

        uz = _mock_contract(country_code='UZ', country_ru='Узбекистан', country_tk='Özbegistan')
        c = ctx.build_contract_context(uz, 'ru')
        self.assertEqual(c['dest_country_gen_ru'], 'Узбекистана')
        self.assertEqual(c['dest_country_gen_tk'], 'Özbegistanyň')
        # The nominative preamble field stays the Country row's own name.
        self.assertEqual(c['buyer_country_ru'], 'Узбекистан')

    def test_unsupported_country_has_no_genitive(self):
        # The view rejects these before rendering; the builder just yields ''.
        tr = _mock_contract(country_code='TR', country_ru='Турция', country_tk='Türkiýe')
        c = ctx.build_contract_context(tr, 'ru')
        self.assertEqual(c['dest_country_gen_ru'], '')
        self.assertEqual(c['dest_country_gen_tk'], '')

    def test_country_template_supported_flag(self):
        for code in ('KZ', 'KG', 'RU', 'TJ', 'UZ', 'AE', 'BY'):
            self.assertTrue(ctx.country_template_supported(code), code)
        self.assertFalse(ctx.country_template_supported('TR'))
        self.assertFalse(ctx.country_template_supported(None))

    def test_financials_and_amount_in_words(self):
        c = ctx.build_contract_context(_mock_contract(), 'ru')
        self.assertEqual(c['contract_no'], '108/26-YGT-EXP')
        self.assertEqual(c['contract_date'], '18.03.2026')
        self.assertEqual(c['total_sum'], '7 830,00')
        self.assertEqual(c['total_sum_words_tk'], 'ýedi müň sekiz ýüz otuz')
        self.assertEqual(c['total_sum_words_ru'], 'семь тысяч восемьсот тридцать')
        self.assertEqual(c['quantity'], '9 000')
        self.assertEqual(c['price'], '0,87')

    def test_dates_spelled_ru_and_tk(self):
        c = ctx.build_contract_context(
            _mock_contract(), 'ru', {'delivery_deadline': '2026-06-30'},
        )
        # RU genitive
        self.assertEqual(c['delivery_deadline_ru'], '30 июня 2026')
        self.assertEqual(c['validity_ru'], '31 декабря 2026')   # from end_date
        # TK ordinal (vowel-harmony suffix): 30 -> -njy, 31 -> -nji
        self.assertEqual(c['delivery_deadline_tk'], '2026-njy ýylyň 30-njy iýunyna')
        self.assertEqual(c['validity_tk'], '2026-njy ýylyň 31-nji dekabryna')

    def test_tk_ordinal_vowel_harmony(self):
        # back-vowel last word -> -njy; thin/front -> -nji
        self.assertEqual([n for n in range(1, 32) if ctx._tk_ordinal(n).endswith('-njy')],
                         [6, 9, 10, 16, 19, 26, 29, 30])

    def test_seller_from_export_firm(self):
        c = ctx.build_contract_context(_mock_contract(), 'ru')
        # legal-form suffix stripped (template supplies "HJ" / "Хозяйственное общество")
        self.assertEqual(c['seller_name_tk'], 'Hemsaya')
        self.assertEqual(c['seller_name_ru'], 'Хемсая')
        # director title word stripped; RU/Cyrillic from `director`, TK/Latin from `director_tk`
        self.assertEqual(c['seller_director_ru'], 'Худайназаров Ы.')
        self.assertEqual(c['seller_director_tk'], 'Hudaýnazarow Ý.')

    def test_seller_director_tk_falls_back_to_ru_when_blank(self):
        c = _mock_contract()
        c.export_firm.director_tk = None  # Turkmen spelling not filled
        out = ctx.build_contract_context(c, 'ru')
        self.assertEqual(out['seller_director_tk'], 'Худайназаров Ы.')  # falls back to `director`
        # bank blob collapsed to one line
        self.assertEqual(out['seller_bank_ru'], 'Банк: Туркменбаши; Вал/счет: 23202; SWIFT: INVATM2X')

    def test_buyer_flat_fields(self):
        c = ctx.build_contract_context(_mock_contract(), 'ru')
        # flat name_company shown in both columns; country is genuinely bilingual
        self.assertEqual(c['buyer_name_tk'], 'TOO «Aranşy - KZ»')
        self.assertEqual(c['buyer_name_ru'], 'TOO «Aranşy - KZ»')
        self.assertEqual(c['buyer_country_tk'], 'Gazagystan')
        self.assertEqual(c['buyer_country_ru'], 'Казахстан')
        # bank_details blob collapses to '; ' (docx runs drop '\n')
        self.assertEqual(
            c['buyer_bank_ru'],
            'БИН 191040016779; БИК HSBKKZKX; р/с KZ97601A891001387241',
        )

    def test_buyer_director_override_wins(self):
        # The modal sends buyer_director (pre-filled from contact_person, editable),
        # so an edited override takes precedence over the firm's stored value.
        c = ctx.build_contract_context(
            _mock_contract(contact_person='Азимов Г.Б.'), 'ru',
            {'buyer_director': 'Edited Name'},
        )
        self.assertEqual(c['buyer_director_tk'], 'Edited Name')
        self.assertEqual(c['buyer_director_ru'], 'Edited Name')

    def test_buyer_director_falls_back_to_contact_person(self):
        # No override, no Turkmen spelling → both columns use contact_person (RU form).
        c = ctx.build_contract_context(_mock_contract(contact_person='Азимов Г.Б.'), 'ru')
        self.assertEqual(c['buyer_director_tk'], 'Азимов Г.Б.')  # falls back to RU form
        self.assertEqual(c['buyer_director_ru'], 'Азимов Г.Б.')

    def test_buyer_director_tk_from_contact_person_tk(self):
        # TK column uses the Latin spelling; RU stays Cyrillic.
        c = ctx.build_contract_context(
            _mock_contract(contact_person='Туктибаев Бекжан',
                           contact_person_tk='Tuktibaýew Bekjan'), 'ru',
        )
        self.assertEqual(c['buyer_director_tk'], 'Tuktibaýew Bekjan')
        self.assertEqual(c['buyer_director_ru'], 'Туктибаев Бекжан')

    def test_blank_deadline_and_director_when_neither_source(self):
        # No override and no start_date → §2.6 stays blank.
        c = ctx.build_contract_context(_mock_contract(start=None), 'ru')
        self.assertEqual(c['delivery_deadline_ru'], '')
        self.assertEqual(c['delivery_deadline_tk'], '')
        self.assertEqual(c['buyer_director_ru'], '')

    def test_header_date_is_contract_date(self):
        # The header ("ş. Asgabat  <date>") carries the document's own date.
        c = ctx.build_contract_context(_mock_contract(contract_date=date(2026, 3, 12)), 'ru')
        self.assertEqual(c['contract_date'], '12.03.2026')

    def test_header_date_falls_back_to_start_date(self):
        # Contracts stored before contract_date existed still print a header date.
        c = ctx.build_contract_context(_mock_contract(), 'ru')
        self.assertEqual(c['contract_date'], '18.03.2026')

    def test_section_2_6_uses_start_date_when_no_override(self):
        c = ctx.build_contract_context(_mock_contract(), 'ru')
        self.assertEqual(c['delivery_deadline_ru'], '18 марта 2026')
        self.assertEqual(c['delivery_deadline_tk'], '2026-njy ýylyň 18-nji martyna')

    def test_section_2_6_override_wins_over_start_date(self):
        c = ctx.build_contract_context(
            _mock_contract(), 'ru', {'delivery_deadline': '2026-06-30'},
        )
        self.assertEqual(c['delivery_deadline_ru'], '30 июня 2026')

    def test_stored_price_per_kg_wins_over_derived(self):
        # 7830 / 9000 = 0,87 derived; the agreed price is what the contract prints.
        c = ctx.build_contract_context(_mock_contract(price='0.9000'), 'ru')
        self.assertEqual(c['price'], '0,9')

    def test_null_financials_render_blank_words(self):
        c = ctx.build_contract_context(_mock_contract(amount=None, qty=None), 'ru')
        self.assertEqual(c['total_sum'], '')
        self.assertEqual(c['total_sum_words_ru'], '')
        self.assertEqual(c['price'], '')

    def test_stamps_off_by_default(self):
        # No ?stamps → all stamp slots blank, even when the firm has images.
        c = _mock_contract()
        c.export_firm.director_seal = SimpleNamespace(name='s.png')
        out = ctx.build_contract_context(c, 'ru')
        self.assertEqual(out['seller_seal'], '')
        self.assertEqual(out['buyer_seal'], '')

    def test_stamps_on_emits_marker_only_when_image_present(self):
        c = _mock_contract()
        c.export_firm.director_seal = SimpleNamespace(name='seal.png')     # uploaded
        c.export_firm.director_signature = None                            # not uploaded
        out = ctx.build_contract_context(c, 'ru', {'stamps': '1'})
        # seller seal → marker; seller signature (no file) → blank
        self.assertIsInstance(out['seller_seal'], ctx.StampImage)
        self.assertEqual(out['seller_signature'], '')
        # buyer has no images in the mock → blank
        self.assertEqual(out['buyer_seal'], '')

    def test_stamps_flag_truthiness(self):
        c = _mock_contract()
        c.export_firm.director_seal = SimpleNamespace(name='seal.png')
        self.assertEqual(ctx.build_contract_context(c, 'ru', {'stamps': '0'})['seller_seal'], '')
        self.assertIsInstance(
            ctx.build_contract_context(c, 'ru', {'stamps': 'true'})['seller_seal'], ctx.StampImage,
        )


class ContractRenderSmokeTest(TestCase):
    """Fill the shipped contract template and assert clean, value-bearing output."""

    def _text(self, data: bytes) -> str:
        doc = Document(BytesIO(data))
        parts = [p.text for p in doc.paragraphs]
        for t in doc.tables:
            for row in t.rows:
                parts.append(' | '.join(c.text for c in row.cells))
        return '\n'.join(parts)

    def test_render_contract_docx(self):
        data, filename, content_type = render.generate(
            'contract_kz', _mock_contract(), 'docx',
            {'buyer_director': 'Tuktibaýew Bekjan', 'delivery_deadline': '2026-06-30'},
        )
        text = self._text(data)
        self.assertNotIn('{{', text)
        self.assertNotIn('{%', text)
        self.assertIn('7 830,00', text)
        self.assertIn('ýedi müň sekiz ýüz otuz', text)      # TK amount in words
        self.assertIn('семь тысяч восемьсот тридцать', text)  # RU amount in words
        self.assertIn('2026-njy ýylyň 30-njy iýunyna', text)  # TK spelled deadline
        self.assertIn('30 июня 2026', text)                   # RU spelled deadline
        self.assertIn('Hemsaya', text)                       # seller name (not hardcoded Ýigit)
        self.assertIn('Худайназаров', text)                  # seller director
        self.assertIn('191040016779', text)                  # buyer bank blob (collapsed)
        self.assertNotIn('Ýigit', text)                      # no leftover hardcoded seller
        self.assertEqual(content_type, render.DOCX_CONTENT_TYPE)
        self.assertEqual(filename, 'Contract_108-26-YGT-EXP_KZ.docx')


class HighlightRenderTest(TestCase):
    """Red-fill of database-driven values (``document_highlight``).

    These assertions deliberately go below ``paragraph.text``: that property
    discards run properties, so the older smoke tests here cannot see colour at
    all, and ``assertNotIn('{{', text)`` still passes on structurally corrupt
    OOXML. The XML-level checks are what actually guard this feature.
    """

    def _runs(self, data: bytes) -> list:
        """Every run in the rendered document, including table cells."""
        doc = Document(BytesIO(data))
        paragraphs = list(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)
        return [run for p in paragraphs for run in p.runs]

    def _red(self, data: bytes) -> list[str]:
        return [
            r.text for r in self._runs(data)
            if r.font.color is not None and r.font.color.rgb == highlight.HIGHLIGHT_RGB
        ]

    def _black(self, data: bytes) -> list[str]:
        return [
            r.text for r in self._runs(data)
            if r.text.strip()
            and (r.font.color is None or r.font.color.rgb is None)
        ]

    def _xml_parts(self, data: bytes) -> dict[str, str]:
        archive = zipfile.ZipFile(BytesIO(data))
        return {
            name: archive.read(name).decode('utf-8', 'ignore')
            for name in archive.namelist() if name.endswith('.xml')
        }

    def test_no_nested_runs_in_any_part(self):
        """A run spliced inside a <w:t> is invalid OOXML that python-docx still parses.

        This is the failure mode of the docxtpl RichText approach, and the reason
        the existing 'no leftover {{' assertion is not sufficient on its own.
        """
        for key, obj in (('invoice_ru', _mock_invoice()),
                         ('contract_kz', _mock_contract())):
            data, _f, _ct = render.generate(key, obj, 'docx')
            for name, xml in self._xml_parts(data).items():
                self.assertIsNone(
                    re.search(r'<w:t[^>]*>[^<]*<w:r>', xml),
                    f'{key}: nested run in {name} — corrupt OOXML',
                )

    def test_no_sentinel_leaks_in_any_part(self):
        """A sentinel reaching the output prints as a tofu box on paper.

        Checked across every XML part, not just word/document.xml — a leak into a
        header would be invisible to a body-only assertion.
        """
        for key, obj in (('invoice_ru', _mock_invoice()),
                         ('contract_kz', _mock_contract())):
            data, _f, _ct = render.generate(key, obj, 'docx')
            for name, xml in self._xml_parts(data).items():
                self.assertNotIn(highlight.OPEN, xml, f'{key}: sentinel in {name}')
                self.assertNotIn(highlight.CLOSE, xml, f'{key}: sentinel in {name}')

    def test_dynamic_values_are_red(self):
        data, _f, _ct = render.generate('invoice_ru', _mock_invoice(), 'docx')
        red = self._red(data)
        self.assertIn('118', red)          # box count, from the shipment
        self.assertIn('7 830,00', red)     # computed total

    def test_boilerplate_stays_black(self):
        """The negative half — over-colouring is as wrong as under-colouring."""
        data, _f, _ct = render.generate('invoice_ru', _mock_invoice(), 'docx')
        black = self._black(data)
        self.assertIn('ИНВОЙС (счет фактура)', black)
        self.assertIn('ПРОДАВЕЦ:', black)
        self.assertNotIn('ИНВОЙС (счет фактура)', self._red(data))

    def test_template_run_formatting_survives_the_split(self):
        """Splitting a run must inherit its rPr, not reset it.

        This is what keeps the invoice's 8/9/10pt hierarchy and the CMR's print
        registration intact.
        """
        data, _f, _ct = render.generate('invoice_ru', _mock_invoice(), 'docx')
        red_runs = [
            r for r in self._runs(data)
            if r.font.color is not None and r.font.color.rgb == highlight.HIGHLIGHT_RGB
        ]
        self.assertTrue(any(r.bold for r in red_runs), 'no bold survived the split')
        self.assertTrue(any(r.font.size is not None for r in red_runs),
                        'no explicit font size survived the split')

    def test_highlight_false_produces_no_colour_and_same_text(self):
        on, _f, _ct = render.generate('invoice_ru', _mock_invoice(), 'docx')
        off, _f2, _ct2 = render.generate(
            'invoice_ru', _mock_invoice(), 'docx', highlight=False,
        )
        self.assertEqual(self._red(off), [])
        self.assertEqual(
            Document(BytesIO(on)).paragraphs[1].text,
            Document(BytesIO(off)).paragraphs[1].text,
            'highlighting must not change the text, only its colour',
        )
        for xml in self._xml_parts(off).values():
            self.assertNotIn(highlight.OPEN, xml)

    def test_blank_values_produce_no_red_run(self):
        """Blank strings are left unwrapped on purpose.

        A wrapped empty value would emit a stray red run, and an unwrapped blank
        stays falsy so a ``{% if %}`` added to a template later still behaves.
        """
        self.assertEqual(highlight.wrap_context({'a': '', 'b': '   '}),
                         {'a': '', 'b': '   '})
        wrapped = highlight.wrap_context({'a': 'x', 'rows': [{'b': 'y', 'c': ''}]})
        self.assertEqual(wrapped['a'], f'{highlight.OPEN}x{highlight.CLOSE}')
        self.assertEqual(wrapped['rows'][0]['b'], f'{highlight.OPEN}y{highlight.CLOSE}')
        self.assertEqual(wrapped['rows'][0]['c'], '')
        self.assertNotIn('', self._red(
            render.generate('invoice_ru', _mock_invoice(), 'docx')[0]))

    def test_non_string_context_values_pass_through(self):
        """Resolved stamps (InlineImage) must never be sentinel-wrapped."""
        marker = object()
        self.assertIs(highlight.wrap_context({'stamp': marker})['stamp'], marker)

    def test_render_does_not_add_document_parts(self):
        """Walking headers/footers must not materialise empty ones into the output."""
        for key, obj in (('invoice_ru', _mock_invoice()),
                         ('contract_kz', _mock_contract())):
            spec = get_spec(key)
            before = set(zipfile.ZipFile(spec.template_path).namelist())
            data, _f, _ct = render.generate(key, obj, 'docx')
            after = set(zipfile.ZipFile(BytesIO(data)).namelist())
            self.assertEqual(after - before, set(), f'{key}: render added parts')

    def test_xlsx_overlay_cells_are_red(self):
        """The CMR overlay colours cells directly — every written cell is dynamic."""
        data, _f, _ct = render.generate('cmr_ru', _mock_shipment(), 'docx')
        ws = openpyxl.load_workbook(BytesIO(data)).active
        filled = [
            c for row in ws.iter_rows() for c in row
            if c.value not in (None, '') and c.font.color is not None
            and getattr(c.font.color, 'rgb', None) == 'FFC00000'
        ]
        self.assertTrue(filled, 'no red cells in the xlsx overlay')

    def test_xlsx_overlay_uncoloured_when_highlight_off(self):
        data, _f, _ct = render.generate(
            'cmr_ru', _mock_shipment(), 'docx', highlight=False,
        )
        ws = openpyxl.load_workbook(BytesIO(data)).active
        reds = [
            c.coordinate for row in ws.iter_rows() for c in row
            if c.font.color is not None
            and getattr(c.font.color, 'rgb', None) == 'FFC00000'
        ]
        self.assertEqual(reds, [])


# Word's storage grids: <w:pgMar> is in twips, <w:sz> in whole half-points.
TWIP_EMU = 635
HALF_POINT_EMU = 6350


class DocumentLayoutRenderTest(TestCase):
    """Saved page-layout adjustments (``DocumentLayoutSetting``) applied at render.

    Deltas and scales, not absolutes — see the model docstring for why.
    """

    def _layout(self, key='invoice_ru', **fields):
        return DocumentLayoutSetting.objects.create(document_key=key, **fields)

    def _doc(self, key='invoice_ru', obj=None):
        data, _f, _ct = render.generate(key, obj or _mock_invoice(), 'docx')
        return Document(BytesIO(data))

    def _runs(self, doc):
        paragraphs = list(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.extend(cell.paragraphs)
        return [r for p in paragraphs for r in p.runs]

    def _template_section(self, key='invoice_ru', index=0):
        return Document(str(get_spec(key).template_path)).sections[index]

    def assertMarginEqual(self, actual, expected, msg=None):
        """Compare margins to within one twip.

        Word stores <w:pgMar> in twentieths of a point, so a margin round-trips
        quantised to 635 EMU. Asserting exact EMU would fail on that rounding
        alone, not on anything the code got wrong.
        """
        self.assertAlmostEqual(actual, expected, delta=TWIP_EMU, msg=msg)

    def test_no_row_leaves_the_template_untouched(self):
        self.assertMarginEqual(
            self._doc().sections[0].left_margin, self._template_section().left_margin,
        )

    def test_default_row_is_a_no_op(self):
        self._layout()
        self.assertMarginEqual(
            self._doc().sections[0].left_margin, self._template_section().left_margin,
        )

    def test_margin_deltas_move_every_section(self):
        template = self._template_section()
        self._layout(margin_left_delta_mm=5, margin_top_delta_mm=-3)
        section = self._doc().sections[0]
        self.assertMarginEqual(section.left_margin, template.left_margin + Mm(5))
        self.assertMarginEqual(section.top_margin, template.top_margin + Mm(-3))

    def test_margin_delta_never_goes_negative(self):
        """contract_kz page 1 has a 0.51cm top margin — a -10mm delta must clamp."""
        self._layout(key='contract_kz', margin_top_delta_mm=-10)
        doc = self._doc('contract_kz', _mock_contract())
        self.assertGreaterEqual(doc.sections[0].top_margin, 0)

    def test_multi_section_margin_difference_is_preserved(self):
        """contract_kz's two sections differ on purpose (letterhead on page 1)."""
        before = Document(str(get_spec('contract_kz').template_path))
        gap = before.sections[1].top_margin - before.sections[0].top_margin

        self._layout(key='contract_kz', margin_top_delta_mm=4)
        after = self._doc('contract_kz', _mock_contract())
        self.assertEqual(after.sections[1].top_margin - after.sections[0].top_margin, gap)

    def test_font_scale_scales_every_sized_run(self):
        """Normal alone is nearly a no-op — most runs carry an explicit <w:sz>."""
        before = {r.text: r.font.size for r in self._runs(self._doc()) if r.font.size}
        self._layout(font_scale_pct=80)
        after = {r.text: r.font.size for r in self._runs(self._doc()) if r.font.size}

        shared = set(before) & set(after)
        self.assertTrue(shared, 'no comparable runs')
        for text in shared:
            # Word stores <w:sz> in whole half-points, so the scaled size snaps to
            # that grid — assert the snap, not a raw EMU product.
            expected = round(before[text] * 0.8 / HALF_POINT_EMU) * HALF_POINT_EMU
            self.assertEqual(after[text], expected, f'run {text!r}')

    def test_line_spacing_reaches_the_document(self):
        self._layout(line_spacing=Decimal('1.50'))
        self.assertEqual(
            self._doc().styles['Normal'].paragraph_format.line_spacing, 1.5,
        )

    def test_fixed_width_tables_track_the_content_width(self):
        """Pinned widths sum to the text area; widen a margin and they must follow.

        Asserts the *ratio*, not an absolute fit — contract_kz's hand-authored
        tables legitimately start wider than their section, and squashing them to
        fit would be us silently redesigning a legal document.
        """
        before = self._doc()
        widths_before = [
            sum(c.width for c in t.rows[0].cells if c.width)
            for t in before.tables if not t.autofit
        ]
        section = before.sections[0]
        content_before = section.page_width - section.left_margin - section.right_margin

        self._layout(margin_left_delta_mm=5, margin_right_delta_mm=5)
        after = self._doc()
        widths_after = [
            sum(c.width for c in t.rows[0].cells if c.width)
            for t in after.tables if not t.autofit
        ]
        section = after.sections[0]
        content_after = section.page_width - section.left_margin - section.right_margin

        self.assertLess(content_after, content_before)   # margins widened
        ratio = content_after / content_before
        self.assertEqual(len(widths_before), len(widths_after))
        for was, now in zip(widths_before, widths_after):
            self.assertAlmostEqual(now / was, ratio, places=3)

    def test_multi_section_document_scales_its_tables_only_once(self):
        """Regression: the rescale used to run per section, squaring the ratio.

        `doc.tables` spans the whole document, so calling it inside the section
        loop applied the ratio twice on contract_kz's two sections (0.8849 where
        0.9402 was correct — a 6% over-shrink). invoice_ru has one section, so
        the single-section test above could never see it.
        """
        template = Document(str(get_spec('contract_kz').template_path))
        widths_before = [
            sum(c.width for c in t.rows[0].cells if c.width)
            for t in template.tables if not t.autofit
        ]
        section = template.sections[0]
        content_before = section.page_width - section.left_margin - section.right_margin

        self._layout(
            key='contract_kz', margin_left_delta_mm=5, margin_right_delta_mm=5,
        )
        after_doc = self._doc('contract_kz', _mock_contract())
        widths_after = [
            sum(c.width for c in t.rows[0].cells if c.width)
            for t in after_doc.tables if not t.autofit
        ]
        section = after_doc.sections[0]
        content_after = section.page_width - section.left_margin - section.right_margin

        self.assertGreater(len(after_doc.sections), 1, 'fixture must be multi-section')
        ratio = content_after / content_before
        for was, now in zip(widths_before, widths_after):
            self.assertAlmostEqual(now / was, ratio, places=3)

    def test_reset_to_defaults_restores_the_pristine_template(self):
        """The way out when an operator has made the layout worse."""
        row = self._layout(font_scale_pct=85, margin_left_delta_mm=8,
                           line_spacing=Decimal('1.40'))
        tuned = self._doc()
        self.assertNotEqual(tuned.sections[0].left_margin,
                            self._template_section().left_margin)

        row.font_scale_pct = 100
        row.line_spacing = None
        row.margin_left_delta_mm = 0
        row.save()

        self.assertTrue(row.is_default)
        reset = self._doc()
        self.assertMarginEqual(
            reset.sections[0].left_margin, self._template_section().left_margin,
        )
        self.assertEqual(
            reset.styles['Normal'].paragraph_format.line_spacing,
            Document(str(get_spec('invoice_ru').template_path))
            .styles['Normal'].paragraph_format.line_spacing,
        )

    def test_autofit_tables_are_left_alone(self):
        """Word reflows autofit tables itself; rescaling them would double-count."""
        before = [
            [c.width for c in t.rows[0].cells]
            for t in self._doc().tables if t.autofit
        ]
        self._layout(margin_left_delta_mm=5)
        after = [
            [c.width for c in t.rows[0].cells]
            for t in self._doc().tables if t.autofit
        ]
        self.assertEqual(before, after)

    def test_cmr_keys_refuse_layout(self):
        """The CMR registers onto pre-printed paper — its geometry is not tunable."""
        for key in ('cmr_ru', 'cmr_en', 'cmr_ru_docx', 'cmr_en_docx'):
            self.assertFalse(tpl_registry.supports_layout(key), key)
            DocumentLayoutSetting.objects.create(
                document_key=key, margin_left_delta_mm=10,
            )
            self.assertIsNone(render.layout_for(key), f'{key} must ignore any row')

    def test_layout_capable_keys_are_the_expected_six(self):
        self.assertEqual(
            set(tpl_registry.layout_capable_keys()),
            {'invoice_ru', 'invoice_en', 'ct1_ru', 'fito_ru', 'customs_tk', 'contract_kz'},
        )

    def test_highlight_and_layout_compose(self):
        """Both post-render passes run; neither undoes the other."""
        self._layout(font_scale_pct=90, margin_left_delta_mm=4)
        data, _f, _ct = render.generate('invoice_ru', _mock_invoice(), 'docx')
        doc = Document(BytesIO(data))
        self.assertMarginEqual(
            doc.sections[0].left_margin, self._template_section().left_margin + Mm(4),
        )
        red = [
            r.text for r in self._runs(doc)
            if r.font.color is not None and r.font.color.rgb == highlight.HIGHLIGHT_RGB
        ]
        self.assertIn('118', red)


class DocumentLayoutModelTest(TestCase):
    """Range validation and the version counter."""

    def test_out_of_range_values_are_rejected(self):
        for field, value in (
            ('font_scale_pct', 130),
            ('line_spacing', Decimal('2.50')),
            ('margin_left_delta_mm', 40),
            ('margin_top_delta_mm', -20),
        ):
            row = DocumentLayoutSetting(document_key='invoice_ru', **{field: value})
            with self.assertRaises(ValidationError, msg=field):
                row.full_clean()

    def test_version_increments_on_update_only(self):
        row = DocumentLayoutSetting.objects.create(document_key='invoice_ru')
        self.assertEqual(row.version, 1)
        row.font_scale_pct = 95
        row.save()
        self.assertEqual(row.version, 2)

    def test_is_default_detects_a_no_op_row(self):
        row = DocumentLayoutSetting(document_key='invoice_ru')
        self.assertTrue(row.is_default)
        row.margin_top_delta_mm = 1
        self.assertFalse(row.is_default)


class InvoiceDocumentEndpointTest(_SeededPermsMixin, TestCase):
    """API: GET /api/v1/contracts/sales/{id}/document/."""

    def setUp(self) -> None:
        self.client = APIClient()
        self.user = _make_user('inv_doc', 'export_manager')
        self.client.force_authenticate(user=self.user)
        self.season = _make_season()
        self.ef = _make_export_firm('YGTDOC')
        self.imp = _make_import_firm('IMPDOC')
        self.contract = _make_contract('INV-DOC-001', self.ef, self.imp, self.season)
        self.invoice = _make_invoice(self.contract, invoice_number=1)
        # Documents require complete shipment packing; link a fully-packed shipment.
        self.invoice.shipment = _make_packed_shipment(self.season, self.imp)
        self.invoice.save(update_fields=['shipment'])

    def test_default_docx_download(self):
        resp = self.client.get(f'/api/v1/contracts/sales/{self.invoice.pk}/document/')
        self.assertEqual(resp.status_code, 200, resp.content[:200])
        self.assertEqual(resp['Content-Type'], render.DOCX_CONTENT_TYPE)
        self.assertIn('attachment;', resp['Content-Disposition'])
        self.assertIn('.docx', resp['Content-Disposition'])
        self.assertGreater(len(resp.content), 1000)

    def test_invoice_en_type(self):
        resp = self.client.get(
            f'/api/v1/contracts/sales/{self.invoice.pk}/document/?type=invoice_en'
        )
        self.assertEqual(resp.status_code, 200)
        self.assertIn('_EN.docx', resp['Content-Disposition'])

    def test_cmr_type_rejected_on_invoice_endpoint(self):
        # CMR is now truck-level (shipment scope); the per-invoice endpoint rejects it.
        resp = self.client.get(
            f'/api/v1/contracts/sales/{self.invoice.pk}/document/?type=cmr_ru'
        )
        self.assertEqual(resp.status_code, 400)

    def test_ct1_letter_type(self):
        resp = self.client.get(
            f'/api/v1/contracts/sales/{self.invoice.pk}/document/?type=ct1_ru'
        )
        self.assertEqual(resp.status_code, 200)
        self.assertIn('CT1_', resp['Content-Disposition'])

    def test_incomplete_packing_returns_400(self):
        # clear a required packing cell → generation is blocked with a clear error
        self.invoice.shipment.pallet_count = None
        self.invoice.shipment.save(update_fields=['pallet_count'])
        resp = self.client.get(f'/api/v1/contracts/sales/{self.invoice.pk}/document/')
        self.assertEqual(resp.status_code, 400)
        self.assertIn('pallet_count', resp.json()['missing_packing'])

    def test_no_shipment_returns_400(self):
        invoice = _make_invoice(self.contract, invoice_number=2)  # no shipment linked
        resp = self.client.get(f'/api/v1/contracts/sales/{invoice.pk}/document/')
        self.assertEqual(resp.status_code, 400)
        self.assertEqual(len(resp.json()['missing_packing']), 4)

    def test_unknown_type_returns_400(self):
        resp = self.client.get(
            f'/api/v1/contracts/sales/{self.invoice.pk}/document/?type=bogus'
        )
        self.assertEqual(resp.status_code, 400)

    def test_pdf_without_libreoffice_returns_503(self):
        with mock.patch.object(render, '_libreoffice_bin', return_value=None):
            resp = self.client.get(
                f'/api/v1/contracts/sales/{self.invoice.pk}/document/?fmt=pdf'
            )
        self.assertEqual(resp.status_code, 503)


class ShipmentCmrEndpointTest(_SeededPermsMixin, TestCase):
    """API: GET /api/v1/contracts/shipments/{id}/cmr/ — truck-level CMR."""

    def setUp(self) -> None:
        self.client = APIClient()
        self.user = _make_user('cmr_doc', 'export_manager')
        self.client.force_authenticate(user=self.user)
        self.season = _make_season()
        self.imp = _make_import_firm('IMPCMR')
        self.shipment = _make_packed_shipment(self.season, self.imp)
        # Two export firms on the one truck → both are senders on the CMR.
        for code in ('YGTA', 'YGTB'):
            ShipmentFirmSplit.objects.create(
                shipment=self.shipment, export_firm=_make_export_firm(code),
                weight_kg=Decimal('9000'), amount_usd=Decimal('8000'),
            )

    def test_truck_cmr_defaults_to_word(self):
        # The office's Word form is the default CMR output.
        resp = self.client.get(f'/api/v1/contracts/shipments/{self.shipment.pk}/cmr/')
        self.assertEqual(resp.status_code, 200, resp.content[:200])
        self.assertEqual(resp['Content-Type'], render.DOCX_CONTENT_TYPE)
        self.assertIn('CMR_', resp['Content-Disposition'])

    def test_en_lang_filename(self):
        resp = self.client.get(
            f'/api/v1/contracts/shipments/{self.shipment.pk}/cmr/?lang=en'
        )
        self.assertEqual(resp.status_code, 200)
        self.assertIn('_EN.docx', resp['Content-Disposition'])

    def test_xlsx_overlay_still_served(self):
        """The spreadsheet overlay is off the UI but still reachable via fmt=xlsx."""
        resp = self.client.get(
            f'/api/v1/contracts/shipments/{self.shipment.pk}/cmr/?fmt=xlsx'
        )
        self.assertEqual(resp.status_code, 200, resp.content[:200])
        self.assertEqual(resp['Content-Type'], render.XLSX_CONTENT_TYPE)
        self.assertIn('.xlsx', resp['Content-Disposition'])

    def test_word_variant(self):
        """fmt=docx returns the editable Word overlay (same values, .docx)."""
        for lang, badge in (('ru', '_RU.docx'), ('en', '_EN.docx')):
            resp = self.client.get(
                f'/api/v1/contracts/shipments/{self.shipment.pk}/cmr/?lang={lang}&fmt=docx'
            )
            self.assertEqual(resp.status_code, 200, resp.content[:200])
            self.assertEqual(resp['Content-Type'], render.DOCX_CONTENT_TYPE)
            self.assertIn(badge, resp['Content-Disposition'])
            text = '\n'.join(
                cell.text
                for table in Document(BytesIO(resp.content)).tables
                for row in table.rows
                for cell in row.cells
            )
            self.assertNotIn('{{', text, f'{lang}: unrendered tag')

    def test_incomplete_packing_returns_400(self):
        self.shipment.box_count = None
        self.shipment.save(update_fields=['box_count'])
        resp = self.client.get(f'/api/v1/contracts/shipments/{self.shipment.pk}/cmr/')
        self.assertEqual(resp.status_code, 400)
        self.assertIn('box_count', resp.json()['missing_packing'])

    def test_missing_shipment_returns_404(self):
        resp = self.client.get('/api/v1/contracts/shipments/999999/cmr/')
        self.assertEqual(resp.status_code, 404)


class ShipmentPacketZipEndpointTest(_SeededPermsMixin, TestCase):
    """API: GET /api/v1/contracts/shipments/{id}/packet.zip — whole packet zip."""

    def setUp(self) -> None:
        self.client = APIClient()
        self.user = _make_user('zip_doc', 'export_manager')
        self.client.force_authenticate(user=self.user)
        self.season = _make_season()
        self.imp = _make_import_firm('IMPZIP')
        self.ef = _make_export_firm('ZIPA')
        self.shipment = _make_packed_shipment(self.season, self.imp, code='0202001/25')
        ShipmentFirmSplit.objects.create(
            shipment=self.shipment, export_firm=self.ef,
            weight_kg=Decimal('9000'), amount_usd=Decimal('8000'),
        )
        contract = _make_contract('ZIP-C1', self.ef, self.imp, self.season)
        sale = _make_invoice(contract, invoice_number=1)
        sale.shipment = self.shipment
        sale.export_firm = self.ef
        sale.save(update_fields=['shipment', 'export_firm'])

    def test_zip_bundles_cmr_invoice_and_letters(self):
        import zipfile
        resp = self.client.get(f'/api/v1/contracts/shipments/{self.shipment.pk}/packet.zip')
        self.assertEqual(resp.status_code, 200, resp.content[:200])
        self.assertEqual(resp['Content-Type'], 'application/zip')
        self.assertIn('.zip', resp['Content-Disposition'])
        names = zipfile.ZipFile(BytesIO(resp.content)).namelist()
        # 1 truck CMR + the firm's invoice + CT-1 + FITO + customs = 5 files.
        self.assertEqual(len(names), 5, names)
        self.assertTrue(any(n.startswith('CMR_') for n in names), names)
        self.assertTrue(any(n.startswith('Invoice_') for n in names), names)

    def test_void_sale_excluded(self):
        import zipfile
        from apps.contracts.models import ContractSale
        ContractSale.objects.filter(shipment=self.shipment).update(status=ContractSale.STATUS_VOID)
        resp = self.client.get(f'/api/v1/contracts/shipments/{self.shipment.pk}/packet.zip')
        self.assertEqual(resp.status_code, 200, resp.content[:200])
        names = zipfile.ZipFile(BytesIO(resp.content)).namelist()
        # voided sale → no invoice/letters; only the truck CMR remains
        self.assertEqual(len(names), 1, names)
        self.assertTrue(names[0].startswith('CMR_'), names)

    def test_incomplete_packing_returns_400(self):
        self.shipment.box_count = None
        self.shipment.save(update_fields=['box_count'])
        resp = self.client.get(f'/api/v1/contracts/shipments/{self.shipment.pk}/packet.zip')
        self.assertEqual(resp.status_code, 400)

    def test_missing_shipment_returns_404(self):
        resp = self.client.get('/api/v1/contracts/shipments/999999/packet.zip')
        self.assertEqual(resp.status_code, 404)


class ContractSaleLineItemApiTest(_SeededPermsMixin, TestCase):
    """API: invoice line items on POST/PATCH /api/v1/contracts/sales/."""

    def setUp(self) -> None:
        self.client = APIClient()
        self.user = _make_user('li_doc', 'export_manager')
        self.client.force_authenticate(user=self.user)
        self.season = _make_season()
        self.ef = _make_export_firm('LIA')
        self.imp = _make_import_firm('LII')
        self.contract = _make_contract('LI-C1', self.ef, self.imp, self.season)

    def _payload(self, lines, quantity='9000', total='8200'):
        return {
            'contract': self.contract.id, 'export_firm': self.ef.id,
            'import_firm': self.imp.id, 'quantity_kg': quantity, 'total_usd': total,
            'line_items': lines,
        }

    def test_create_with_matching_lines(self):
        lines = [
            {'product_name': 'Сорт А', 'quantity_kg': '5000', 'price_per_kg': '1.0000'},
            {'product_name': 'Сорт Б', 'quantity_kg': '4000', 'price_per_kg': '0.8000'},
        ]
        resp = self.client.post('/api/v1/contracts/sales/', self._payload(lines), format='json')
        self.assertEqual(resp.status_code, 201, resp.content[:300])
        detail = self.client.get(f"/api/v1/contracts/sales/{resp.json()['id']}/").json()
        self.assertEqual(len(detail['line_items']), 2)
        self.assertEqual(detail['line_items'][0]['line_number'], 1)   # server-assigned
        self.assertEqual(detail['line_items'][0]['total_usd'], '5000.00')  # server-computed

    def test_lines_must_sum_to_total(self):
        lines = [
            {'product_name': 'A', 'quantity_kg': '5000', 'price_per_kg': '1.0000'},
            {'product_name': 'B', 'quantity_kg': '4000', 'price_per_kg': '0.5000'},  # → 7000 ≠ 8200
        ]
        resp = self.client.post('/api/v1/contracts/sales/', self._payload(lines), format='json')
        self.assertEqual(resp.status_code, 400)

    def test_lines_must_sum_to_quantity(self):
        lines = [{'product_name': 'A', 'quantity_kg': '5000', 'price_per_kg': '1.6400'}]  # 5000≠9000
        resp = self.client.post('/api/v1/contracts/sales/', self._payload(lines), format='json')
        self.assertEqual(resp.status_code, 400)

    def test_money_patch_revalidates_existing_lines(self):
        # existing lines sum to 8200; PATCHing total_usd away from that (without
        # resending line_items) must be rejected, not leave a stale mismatch.
        create = self.client.post('/api/v1/contracts/sales/', self._payload([
            {'product_name': 'A', 'quantity_kg': '5000', 'price_per_kg': '1.0000'},
            {'product_name': 'B', 'quantity_kg': '4000', 'price_per_kg': '0.8000'},
        ]), format='json')
        sale_id = create.json()['id']
        bad = self.client.patch(
            f'/api/v1/contracts/sales/{sale_id}/', {'total_usd': '9999'}, format='json',
        )
        self.assertEqual(bad.status_code, 400)

    def test_status_only_patch_skips_line_check(self):
        # a status-only PATCH doesn't touch money → no re-reconciliation, no block.
        create = self.client.post('/api/v1/contracts/sales/', self._payload([
            {'product_name': 'A', 'quantity_kg': '9000', 'price_per_kg': '0.9111'},
        ], total='8199.90'), format='json')
        sale_id = create.json()['id']
        ok = self.client.patch(
            f'/api/v1/contracts/sales/{sale_id}/', {'status': 'paid'}, format='json',
        )
        self.assertEqual(ok.status_code, 200, ok.content[:300])

    def test_patch_empty_list_clears_lines(self):
        create = self.client.post('/api/v1/contracts/sales/', self._payload([
            {'product_name': 'A', 'quantity_kg': '5000', 'price_per_kg': '1.0000'},
            {'product_name': 'B', 'quantity_kg': '4000', 'price_per_kg': '0.8000'},
        ]), format='json')
        sale_id = create.json()['id']
        patch = self.client.patch(
            f'/api/v1/contracts/sales/{sale_id}/', {'line_items': []}, format='json',
        )
        self.assertEqual(patch.status_code, 200, patch.content[:300])
        detail = self.client.get(f'/api/v1/contracts/sales/{sale_id}/').json()
        self.assertEqual(detail['line_items'], [])


class ContractAgreementEndpointTest(_SeededPermsMixin, TestCase):
    """API: GET /api/v1/contracts/contracts/{id}/agreement/."""

    def setUp(self) -> None:
        self.client = APIClient()
        self.user = _make_user('ctr_doc', 'export_manager')
        self.client.force_authenticate(user=self.user)
        self.season = _make_season()
        self.ef = _make_export_firm('YGTCTR')
        self.imp = _make_import_firm('IMPCTR')
        self.imp.country = _make_country()  # Kazakhstan (KZ) — required by the gate
        self.imp.save(update_fields=['country'])
        self.contract = _make_contract('108/26-YGT-EXP', self.ef, self.imp, self.season)
        self.contract.planned_amount_usd = Decimal('7830.00')
        self.contract.planned_quantity_kg = Decimal('9000')
        self.contract.start_date = date(2026, 3, 18)
        self.contract.end_date = date(2026, 12, 31)
        self.contract.save()

    def test_default_docx_download(self):
        resp = self.client.get(
            f'/api/v1/contracts/contracts/{self.contract.pk}/agreement/'
            '?buyer_director=Tuktibayew&delivery_deadline=2026-06-30'
        )
        self.assertEqual(resp.status_code, 200, resp.content[:200])
        self.assertEqual(resp['Content-Type'], render.DOCX_CONTENT_TYPE)
        self.assertIn('attachment;', resp['Content-Disposition'])
        self.assertIn('Contract_108-26-YGT-EXP_KZ.docx', resp['Content-Disposition'])
        self.assertGreater(len(resp.content), 1000)

    def test_supported_non_kz_buyer_downloads(self):
        # §4 names the destination country, so any country with a verified genitive
        # form generates — not Kazakhstan alone.
        uz, _ = Country.objects.get_or_create(
            code='UZ', defaults={'name_tk': 'Özbegistan', 'name_ru': 'Узбекистан', 'name_en': 'Uzbekistan'},
        )
        self.imp.country = uz
        self.imp.save(update_fields=['country'])
        resp = self.client.get(f'/api/v1/contracts/contracts/{self.contract.pk}/agreement/')
        self.assertEqual(resp.status_code, 200, resp.content[:200])
        self.assertGreater(len(resp.content), 1000)
        # §4 must name Uzbekistan's authorities, with no Kazakhstan left behind.
        body = zipfile.ZipFile(BytesIO(resp.content)).read('word/document.xml').decode('utf8')
        text = re.sub(r'<[^>]+>', '', body)
        self.assertIn('Узбекистана', text)
        self.assertIn('Özbegistanyň', text)
        self.assertNotIn('Казахстан', text)
        self.assertNotIn('Gazagystany', text)

    def test_unsupported_country_returns_400(self):
        # No verified genitive form → the §4 clauses cannot be worded correctly.
        tr, _ = Country.objects.get_or_create(
            code='TR', defaults={'name_tk': 'Türkiýe', 'name_ru': 'Турция', 'name_en': 'Turkey'},
        )
        self.imp.country = tr
        self.imp.save(update_fields=['country'])
        resp = self.client.get(f'/api/v1/contracts/contracts/{self.contract.pk}/agreement/')
        self.assertEqual(resp.status_code, 400)
        self.assertIn('TR', resp.json()['error'])

    def test_country_not_set_returns_400(self):
        self.imp.country = None
        self.imp.save(update_fields=['country'])
        resp = self.client.get(f'/api/v1/contracts/contracts/{self.contract.pk}/agreement/')
        self.assertEqual(resp.status_code, 400)

    def test_pdf_without_libreoffice_returns_503(self):
        with mock.patch.object(render, '_libreoffice_bin', return_value=None):
            resp = self.client.get(
                f'/api/v1/contracts/contracts/{self.contract.pk}/agreement/?fmt=pdf'
            )
        self.assertEqual(resp.status_code, 503)


class DocumentPacketEndpointTest(_SeededPermsMixin, TestCase):
    """API: GET /api/v1/contracts/document-packets/ — one row per truck."""

    def setUp(self) -> None:
        self.client = APIClient()
        self.user = _make_user('pkt_doc', 'export_manager')
        self.client.force_authenticate(user=self.user)
        self.season = _make_season()
        self.season.is_active = True
        self.season.save(update_fields=['is_active'])
        self.imp = _make_import_firm('IMPPKT')
        self.ef1 = _make_export_firm('PKTA')
        self.ef2 = _make_export_firm('PKTB')
        self.shipment = _make_packed_shipment(
            self.season, self.imp, status_code='yola_chykdy', code='0101701/25',
        )
        for firm in (self.ef1, self.ef2):
            ShipmentFirmSplit.objects.create(
                shipment=self.shipment, export_firm=firm,
                weight_kg=Decimal('9000'), amount_usd=Decimal('8000'),
            )
        # ef1 has a linked sale (invoice); ef2 does not yet.
        contract = _make_contract('PKT-C1', self.ef1, self.imp, self.season)
        sale = _make_invoice(contract, invoice_number=1)
        sale.shipment = self.shipment
        sale.export_firm = self.ef1
        sale.save(update_fields=['shipment', 'export_firm'])

    def test_lists_truck_packet(self):
        resp = self.client.get('/api/v1/contracts/document-packets/')
        self.assertEqual(resp.status_code, 200)
        results = resp.json()['results']
        self.assertEqual(len(results), 1)
        pkt = results[0]
        self.assertEqual(pkt['id'], self.shipment.id)
        self.assertTrue(pkt['packing_complete'])
        self.assertEqual(pkt['buyer_name'], self.imp.name_short or self.imp.name_company)
        self.assertTrue(pkt['is_ready'])           # fully filled → ready
        self.assertEqual(pkt['missing_setup'], [])
        by_firm = {f['export_firm_id']: f for f in pkt['firms']}
        self.assertEqual(len(by_firm), 2)
        self.assertIsNotNone(by_firm[self.ef1.id]['sale_id'])   # ef1 has a sale
        self.assertIsNone(by_firm[self.ef2.id]['sale_id'])      # ef2 does not

    def test_document_team_can_list(self):
        # Regression: document_team is the page's primary user — it must have the
        # 'sale' resource (view) or every Documents-page call 403s.
        doc_user = _make_user('pkt_dt', 'document_team')
        client = APIClient()
        client.force_authenticate(user=doc_user)
        resp = client.get('/api/v1/contracts/document-packets/')
        self.assertEqual(resp.status_code, 200)

    def test_filled_draft_is_included(self):
        # a DRAFT with all the gate fields (firms + buyer + country + driver + plate)
        # now appears — the page gates on data readiness, not lifecycle status.
        draft = _make_packed_shipment(
            self.season, self.imp, status_code='draft', code='0101702/25',
        )
        ShipmentFirmSplit.objects.create(
            shipment=draft, export_firm=self.ef1,
            weight_kg=Decimal('9000'), amount_usd=Decimal('8000'),
        )
        ids = [p['id'] for p in self.client.get('/api/v1/contracts/document-packets/').json()['results']]
        self.assertIn(draft.id, ids)

    def test_incomplete_truck_shown_with_flags(self):
        # missing driver + plate → still SHOWS, flagged not-ready, so the team sees
        # what to fill instead of the truck silently vanishing.
        incomplete = _make_packed_shipment(
            self.season, self.imp, status_code='yola_chykdy', code='0101703/25',
            driver_name='', truck_plate='',
        )
        ShipmentFirmSplit.objects.create(
            shipment=incomplete, export_firm=self.ef1,
            weight_kg=Decimal('9000'), amount_usd=Decimal('8000'),
        )
        by_id = {p['id']: p for p in self.client.get('/api/v1/contracts/document-packets/').json()['results']}
        self.assertIn(incomplete.id, by_id)
        pkt = by_id[incomplete.id]
        self.assertFalse(pkt['is_ready'])
        self.assertIn('driver_name', pkt['missing_setup'])
        self.assertIn('truck_plate', pkt['missing_setup'])

    def test_truck_without_firms_hidden(self):
        # floor: no export firm assigned → nothing to invoice / no CMR sender, hidden
        bare = _make_packed_shipment(self.season, self.imp, status_code='draft', code='0101704/25')
        ids = [p['id'] for p in self.client.get('/api/v1/contracts/document-packets/').json()['results']]
        self.assertNotIn(bare.id, ids)

    def test_incomplete_packing_flag(self):
        self.shipment.box_count = None
        self.shipment.save(update_fields=['box_count'])
        pkt = self.client.get('/api/v1/contracts/document-packets/').json()['results'][0]
        self.assertFalse(pkt['packing_complete'])
        self.assertIn('box_count', pkt['missing_packing'])

    def test_firm_filter(self):
        on = self.client.get(f'/api/v1/contracts/document-packets/?firm={self.ef1.id}')
        self.assertEqual(len(on.json()['results']), 1)
        other = _make_export_firm('PKTC')
        off = self.client.get(f'/api/v1/contracts/document-packets/?firm={other.id}')
        self.assertEqual(len(off.json()['results']), 0)


class DocumentLayoutEndpointTest(_SeededPermsMixin, TestCase):
    """API: /api/v1/contracts/document-layouts/."""

    LIST_URL = '/api/v1/contracts/document-layouts/'

    def setUp(self) -> None:
        self.client = APIClient()
        self.user = _make_user('layout_mgr', 'export_manager')
        self.client.force_authenticate(user=self.user)

    def _detail(self, key: str) -> str:
        return f'{self.LIST_URL}{key}/'

    def test_list_synthesises_defaults_for_untouched_documents(self):
        resp = self.client.get(self.LIST_URL)
        self.assertEqual(resp.status_code, 200, resp.content[:200])
        by_key = {row['document_key']: row for row in resp.json()}
        self.assertEqual(
            set(by_key),
            {'invoice_ru', 'invoice_en', 'ct1_ru', 'fito_ru', 'customs_tk', 'contract_kz'},
        )
        self.assertEqual(by_key['invoice_ru']['font_scale_pct'], 100)
        self.assertEqual(by_key['invoice_ru']['margin_left_delta_mm'], 0)
        self.assertIsNone(by_key['invoice_ru']['line_spacing'])

    def test_list_excludes_the_cmr_keys(self):
        keys = {row['document_key'] for row in self.client.get(self.LIST_URL).json()}
        for locked in ('cmr_ru', 'cmr_en', 'cmr_ru_docx', 'cmr_en_docx'):
            self.assertNotIn(locked, keys)

    def test_patch_creates_the_row_on_first_edit(self):
        resp = self.client.patch(
            self._detail('invoice_ru'), {'font_scale_pct': 90}, format='json',
        )
        self.assertEqual(resp.status_code, 200, resp.content[:200])
        self.assertEqual(resp.json()['font_scale_pct'], 90)
        row = DocumentLayoutSetting.objects.get(document_key='invoice_ru')
        self.assertEqual(row.updated_by, self.user)

    def test_patch_is_partial(self):
        self.client.patch(
            self._detail('invoice_ru'), {'font_scale_pct': 90}, format='json',
        )
        self.client.patch(
            self._detail('invoice_ru'), {'margin_left_delta_mm': 4}, format='json',
        )
        row = DocumentLayoutSetting.objects.get(document_key='invoice_ru')
        self.assertEqual(row.font_scale_pct, 90)     # not reset by the second PATCH
        self.assertEqual(row.margin_left_delta_mm, 4)

    def test_out_of_range_is_rejected(self):
        resp = self.client.patch(
            self._detail('invoice_ru'), {'font_scale_pct': 300}, format='json',
        )
        self.assertEqual(resp.status_code, 400)
        self.assertIn('font_scale_pct', resp.json())

    def test_stale_version_conflicts(self):
        self.client.patch(
            self._detail('invoice_ru'), {'font_scale_pct': 90}, format='json',
        )
        current = DocumentLayoutSetting.objects.get(document_key='invoice_ru').version
        resp = self.client.patch(
            self._detail('invoice_ru'),
            {'font_scale_pct': 95, 'version': current - 1},
            format='json',
        )
        self.assertEqual(resp.status_code, 409)
        self.assertEqual(resp.json()['current_version'], current)
        self.assertEqual(
            DocumentLayoutSetting.objects.get(document_key='invoice_ru').font_scale_pct,
            90,  # unchanged
        )

    def test_matching_version_succeeds(self):
        self.client.patch(
            self._detail('invoice_ru'), {'font_scale_pct': 90}, format='json',
        )
        current = DocumentLayoutSetting.objects.get(document_key='invoice_ru').version
        resp = self.client.patch(
            self._detail('invoice_ru'),
            {'font_scale_pct': 95, 'version': current},
            format='json',
        )
        self.assertEqual(resp.status_code, 200, resp.content[:200])

    def test_reset_clears_line_spacing_back_to_null(self):
        """The popover's Reset PATCHes an explicit null — it must not be dropped.

        `partial=True` ignores absent keys, so an explicit JSON ``null`` is the
        only way back to "use the template's own spacing".
        """
        self.client.patch(
            self._detail('invoice_ru'), {'line_spacing': '1.40'}, format='json',
        )
        resp = self.client.patch(
            self._detail('invoice_ru'),
            {'line_spacing': None, 'font_scale_pct': 100, 'margin_left_delta_mm': 0},
            format='json',
        )
        self.assertEqual(resp.status_code, 200, resp.content[:200])
        self.assertIsNone(resp.json()['line_spacing'])
        row = DocumentLayoutSetting.objects.get(document_key='invoice_ru')
        self.assertIsNone(row.line_spacing)
        self.assertTrue(row.is_default, 'a full reset must render as a no-op')

    def test_cmr_key_is_refused(self):
        resp = self.client.patch(
            self._detail('cmr_ru_docx'), {'margin_left_delta_mm': 5}, format='json',
        )
        self.assertEqual(resp.status_code, 400)
        self.assertFalse(DocumentLayoutSetting.objects.exists())

    def test_unknown_key_is_refused(self):
        resp = self.client.patch(
            self._detail('not_a_document'), {'font_scale_pct': 90}, format='json',
        )
        self.assertEqual(resp.status_code, 400)

    def test_anonymous_cannot_read_or_write(self):
        anon = APIClient()
        self.assertEqual(anon.get(self.LIST_URL).status_code, 401)
        self.assertEqual(
            anon.patch(self._detail('invoice_ru'), {'font_scale_pct': 90},
                       format='json').status_code,
            401,
        )

    def test_saved_layout_reaches_a_generated_document(self):
        """The end-to-end point of the feature: tune it, download it, see it."""
        self.client.patch(
            self._detail('invoice_ru'), {'margin_left_delta_mm': 6}, format='json',
        )
        template_margin = Document(
            str(get_spec('invoice_ru').template_path)
        ).sections[0].left_margin
        data, _f, _ct = render.generate('invoice_ru', _mock_invoice(), 'docx')
        rendered = Document(BytesIO(data)).sections[0].left_margin
        self.assertAlmostEqual(rendered, template_margin + Mm(6), delta=TWIP_EMU)
