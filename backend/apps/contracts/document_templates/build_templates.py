"""Generate the invoice ``.docx`` templates (Jinja-tagged for docxtpl).

Run once to (re)create ``invoice_ru.docx`` / ``invoice_en.docx``:

    python -m apps.contracts.document_templates.build_templates

The produced ``.docx`` files are the source-of-truth layouts committed to git.
Static labels (ИНВОЙС/INVOICE, ПРОДАВЕЦ/SELLER, table headers …) are baked into
each language file; only data values are ``{{ jinja }}`` tags, filled at runtime
by ``apps.contracts.services.document_render``.

NOTE: once a human refines a template in Word, that ``.docx`` is authoritative —
re-running this overwrites it. Only re-run to intentionally reset a template.
"""
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUT_DIR = Path(__file__).resolve().parent

# Per-language static label set. Data values stay as Jinja tags shared by both.
# The invoice .docx mirrors the office Excel sheet (InvoiceRU / InvoiceEN): a first
# INVOICE page and a second PACKING-LIST page ("Упаковочный лист" / "Packing List")
# that repeats the parties/route but drops the price columns.
LABELS = {
    'ru': {
        'title': 'ИНВОЙС (счет фактура)',
        'packing_title': 'Упаковочный лист',
        'number_line': '№ {{ invoice_no }} от   {{ invoice_date }}',
        'contract_line': 'Контракт № {{ contract_line }}',
        'seller': 'ПРОДАВЕЦ:',
        'buyer': 'ПОКУПАТЕЛЬ:',
        'country': 'Страна происхождение товара:',
        'loading': 'Место погрузки груза:',
        'delivery': 'Условие поставки:',
        'transport': 'Вид транспорта: Авто:',
        'cols': ['№', 'Наименование товара', 'Код по ТН ВЭД', 'Кол-во мест',
                 'Род упаковки', 'Брутто', 'Нетто',
                 'Цена долл.США за 1 кг', 'Сумма долл.США'],
        'packing_cols': ['№', 'Наименование товара', 'Код по ТН ВЭД',
                         'Кол-во мест', 'Брутто', 'Нетто'],
        'total': 'ИТОГО:',
        'seller_foot': 'ПРОДАВЕЦ',
        'released': 'Товар отпустил: __________________{{ seller_name }}',
        'sign': '(подпись лица)',
    },
    'en': {
        'title': 'INVOICE',
        'packing_title': 'Packing List',
        'number_line': '№ {{ invoice_no }}, {{ invoice_date }}',
        'contract_line': 'Contract № {{ contract_line }}',
        'seller': 'SELLER:',
        'buyer': 'BUYER:',
        'country': 'Country of origin:',
        'loading': 'Place of loading cargo:',
        'delivery': 'Delivery conditions:',
        'transport': 'Type of transport: Auto:',
        'cols': ['№', 'Name of product', 'Code on TN FEA', 'Number of pieces',
                 'Type of packing', 'Gross weight, kg', 'Net weight, kg',
                 'Price of US dollars kg', 'Total price of US dollars'],
        'packing_cols': ['№', 'Name of product', 'Code on TN FEA',
                         'Number of pieces', 'Gross weight, kg', 'Net weight, kg'],
        'total': 'TOTAL:',
        'seller_foot': 'SELLER',
        'released': ' __________________{{ seller_name }}',
        'sign': '(signature)',
    },
}

# Data field per invoice column (9-col invoice table / 6-col packing table).
_INVOICE_FIELDS = ['n', 'name', 'code', 'pieces', 'packing', 'gross', 'net', 'price', 'total']
_PACKING_FIELDS = ['n', 'name', 'code', 'pieces', 'gross', 'net']


def _set_cell(cell, text, bold=False, size=9, italic=False):
    cell.text = ''
    run = cell.paragraphs[0].add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)


def _no_border(cell) -> None:
    """Strip all borders from a single cell (used for the spacer between boxes)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'nil')
        borders.append(el)
    tc_pr.append(borders)


def _col_widths(table, widths) -> None:
    """Pin column widths (python-docx needs each cell set, autofit off)."""
    table.autofit = False
    table.allow_autofit = False
    for row in table.rows:
        for cell, width in zip(row.cells, widths):
            cell.width = width


def _a4(doc) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2)
    section.right_margin = Cm(1.5)


def _header(doc, lab, title_text: str) -> None:
    """Centered title + № line + contract line (shared by both invoice pages)."""
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(title_text)
    r.bold = True
    r.font.size = Pt(13)
    for key in ('number_line', 'contract_line'):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(lab[key]).bold = True
    doc.add_paragraph()


def _fill_party_cell(cell, name_tag: str, addr_tag: str, bank_tag: str) -> None:
    """Firm box: name (top, bold), address, gap, then multi-line bank block.

    ``bank_tag`` resolves to the firm's bank-details blob whose embedded ``\\n``
    docxtpl renders as line breaks — so the box mirrors the Excel requisites list.
    """
    cell.text = ''
    top = cell.paragraphs[0].add_run(name_tag)
    top.bold = True
    top.font.size = Pt(10)
    cell.add_paragraph().add_run(addr_tag).font.size = Pt(9)
    for _ in range(3):
        cell.add_paragraph()
    cell.add_paragraph().add_run(bank_tag).font.size = Pt(8)


def _party_block(doc, lab) -> None:
    """ПРОДАВЕЦ / ПОКУПАТЕЛЬ labels over two tall bordered firm boxes with a gap."""
    widths = [Cm(8), Cm(1.5), Cm(8)]
    labels = doc.add_table(rows=1, cols=3)
    _col_widths(labels, widths)
    _set_cell(labels.cell(0, 0), lab['seller'], bold=True, size=10)
    _set_cell(labels.cell(0, 2), lab['buyer'], bold=True, size=10)

    boxes = doc.add_table(rows=1, cols=3)
    boxes.style = 'Table Grid'
    _col_widths(boxes, widths)
    boxes.rows[0].height = Cm(8)
    boxes.rows[0].height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    _fill_party_cell(boxes.cell(0, 0), '{{ seller_name }}', '{{ seller_address }}', '{{ seller_bank }}')
    _no_border(boxes.cell(0, 1))
    _fill_party_cell(boxes.cell(0, 2), '{{ buyer_name }}', '{{ buyer_address }}', '{{ buyer_bank }}')


def _info_block(doc, lab) -> None:
    """Route rows: origin / place of loading / delivery terms / transport."""
    rows = (
        (lab['country'], '{{ country_origin }}'),
        (lab['loading'], '{{ place_loading }}'),
        (lab['delivery'], '{{ delivery_terms }}'),
        (lab['transport'], '{{ transport }}'),
    )
    table = doc.add_table(rows=len(rows), cols=2)
    _col_widths(table, [Cm(6), Cm(11.5)])
    for i, (label, tag) in enumerate(rows):
        _set_cell(table.rows[i].cells[0], label, size=10)
        _set_cell(table.rows[i].cells[1], tag, size=10)
    doc.add_paragraph()


def _items_table(doc, cols, fields, total_label=None) -> None:
    """Bordered line-items table: header + a docxtpl row-loop over ``line_items``
    (+ optional ИТОГО row). The data row is emitted once per line item at render
    time — one product line (the common case) or several (varieties/grades).

    docxtpl 0.19 row loop: the ``{%tr for%}`` / ``{%tr endfor%}`` markers must sit
    in their OWN rows (bracketing the data row); docxtpl removes the two marker
    rows and repeats the data row between them. (Both tags in one row — the older
    idiom — is rejected as 'unknown tag endfor' by this version.)
    """
    # header + for-marker + data + endfor-marker (+ optional total)
    n_rows = 4 + (1 if total_label else 0)
    table = doc.add_table(rows=n_rows, cols=len(cols))
    table.style = 'Table Grid'
    for i, name in enumerate(cols):
        _set_cell(table.rows[0].cells[i], name, bold=True, size=8)
    _set_cell(table.rows[1].cells[0], '{%tr for item in line_items %}', size=9)
    data = table.rows[2].cells
    for i, fld in enumerate(fields):
        _set_cell(data[i], f'{{{{ item.{fld} }}}}', size=9)
    _set_cell(table.rows[3].cells[0], '{%tr endfor %}', size=9)
    if total_label:
        total_cells = table.rows[4].cells
        _set_cell(total_cells[1], total_label, bold=True)
        _set_cell(total_cells[len(cols) - 1], '{{ total_sum }}', bold=True)


def _footer(doc, lab) -> None:
    """Pallet note + seller signature block (shared by both invoice pages)."""
    note = doc.add_paragraph()
    note.add_run('{{ pallet_note }}').italic = True
    doc.add_paragraph()
    doc.add_paragraph().add_run(lab['seller_foot']).bold = True
    doc.add_paragraph().add_run(lab['released']).bold = True
    sign = doc.add_paragraph()
    sign.add_run(lab['sign']).italic = True


def build(lang: str) -> Path:
    lab = LABELS[lang]
    doc = Document()
    doc.styles['Normal'].font.size = Pt(10)
    _a4(doc)

    # Page 1 — INVOICE (with price columns and ИТОГО total).
    _header(doc, lab, lab['title'])
    _party_block(doc, lab)
    doc.add_paragraph()
    _info_block(doc, lab)
    _items_table(doc, lab['cols'], _INVOICE_FIELDS, total_label=lab['total'])
    _footer(doc, lab)

    # Page 2 — PACKING LIST (same header/parties/route, weights only, no prices).
    doc.add_page_break()
    _header(doc, lab, lab['packing_title'])
    _party_block(doc, lab)
    doc.add_paragraph()
    _info_block(doc, lab)
    _items_table(doc, lab['packing_cols'], _PACKING_FIELDS)
    _footer(doc, lab)

    out = OUT_DIR / f'invoice_{lang}.docx'
    doc.save(out)
    return out


# NOTE: the CMR is NOT built here — it is an xlsx print-overlay onto the
# pre-printed official form (geometry python-docx can't reproduce). See
# ``build_cmr_xlsx.py`` and ``document_context.build_cmr_overlay``.


# Authority request letters — single-language forms mirroring the office Excel
# sheets (`letter CT1` / `fito` / `customs`). Each is a self-contained builder:
# addressee → body → sender/consignee blocks → weights/table → signature. Static
# boilerplate is baked in; only the named {{ fields }} are injected at render time.

# Letters use a serif face like the office correspondence.
_LETTER_FONT = 'Times New Roman'


def _letter_doc() -> 'Document':
    doc = Document()
    _a4(doc)
    normal = doc.styles['Normal'].font
    normal.name = _LETTER_FONT
    normal.size = Pt(12)
    return doc


def _addressee(doc, lines) -> None:
    """Right-aligned bold addressee block, indented into the right half."""
    for line in lines:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.left_indent = Cm(7)
        p.add_run(line).bold = True


def _para(doc, text, *, justify=False, bold=False, size=12):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    return p


def _weight_rows(doc, rows) -> None:
    """Borderless label/value grid (Нетто / Брутто / Кол-во мест)."""
    table = doc.add_table(rows=len(rows), cols=2)
    _col_widths(table, [Cm(3.5), Cm(6)])
    for i, (label, tag) in enumerate(rows):
        _set_cell(table.rows[i].cells[0], label, size=12)
        _set_cell(table.rows[i].cells[1], tag, size=12)


def build_ct1() -> Path:
    """CT-1 certificate-of-origin request letter (RU)."""
    doc = _letter_doc()
    _addressee(doc, ['Директору предприятия', '«Туркменэкспертиза» ТПП в', 'Туркменистане'])
    doc.add_paragraph()
    _para(doc, '{{ firm_name }} просит Вас оформить сертификат происхождения '
               'формы «СТ-1», на {{ product }}. Контракт № {{ contract_line }}.', justify=True)
    _para(doc, 'Отправитель: {{ firm_name }}')
    _para(doc, '{{ firm_address }}')
    doc.add_paragraph()
    _para(doc, 'Грузополучатель: {{ buyer_name }}')
    _para(doc, '{{ buyer_address }}', size=9)
    doc.add_paragraph()
    _weight_rows(doc, [('Нетто:', '{{ net }} кг.'), ('Брутто:', '{{ gross }} кг.'),
                       ('Кол-во мест:', '{{ boxes }} шт.')])
    doc.add_paragraph()
    _para(doc, '{{ firm_name }}__________________________', bold=True)

    out = OUT_DIR / 'ct1_ru.docx'
    doc.save(out)
    return out


def build_fito() -> Path:
    """Phytosanitary-certificate request letter (RU)."""
    doc = _letter_doc()
    _addressee(doc, ['Гос служба по карантину', 'растений Ахалского велаята'])
    doc.add_paragraph()
    _para(doc, '{{ firm_name }} просит Вас оформить Фитосанитарный сертификат на груз, '
               'направлению в {{ country }}. Вес груза нетто {{ net }} кг., '
               'ящика - {{ boxes }} шт., на {{ product }}.', justify=True)
    doc.add_paragraph()
    _para(doc, '1 автомашина: {{ plate }}')
    doc.add_paragraph()
    _para(doc, 'Отправитель: {{ firm_name }}')
    _para(doc, '{{ firm_address }}')
    _para(doc, 'Грузополучатель: {{ buyer_name }}')
    doc.add_paragraph()
    _para(doc, '{{ buyer_address }}', size=9)
    doc.add_paragraph()
    doc.add_paragraph()
    _para(doc, '{{ firm_name }}__________________________', bold=True)

    out = OUT_DIR / 'fito_ru.docx'
    doc.save(out)
    return out


# Static Turkmen legal boilerplate for the customs ARZA (verbatim from the sheet).
_CUSTOMS_BODY = (
    '{{ seller_name }} bilen {{ buyer_name }} arasynda {{ contract_line }} senede '
    'baglaşylan şertnama boýunça ýükümizi {{ country }} Respublikasyna çykarmak üçin '
    'gümrük gözegçisini bermegiňizi Sizden haýyş edýäris.'
)
_CUSTOMS_PARAS = [
    'Harytlaryň ýüklenjek ýeri: {{ place_loading }} Bu harytlaryň arasynda '
    'Türkmenistanyň çäginden alnyp gidilmegi gadagan edilen zatlaryň we neşe '
    'serişdeleriniň ýokdugyna güwa geçmek bilen, Türkmenistanyň Kanunçylygynda '
    'bellenen tertipde we möhletde degişli gümrük töleglerini wagtynda tölemäge '
    'hem-de 10 günüň dowamynda görkezilen harytlary awtoulag serişdesine (demirýol '
    'wagonlaryna) ýüklemäge we olary gümrük taýdan resmileşdirmek üçin ÝGD-ny, '
    'gümrük edarasyna berilmegi göz öňüne tutulan beýleki resminamalary gümrük '
    'edarasyna eltip bermäge borçlanýarys',
    'Türkmenistanyň gümrük Kodeksiniň 31,47,53,77,78,81,273 Türkmenistanyň '
    'administratiw-hukuk tertibiniň bozulmalary hakyndaky kodeksiniň 390-407-nji '
    'hem-de jenaýat kodeksiniň 261-nji maddalary we bu düzgünleriň bozulmagy üçin '
    'jogapkärçilik babatda doly düşündirildi.',
    'Türkmenistanyň Maliýe we ykdysadyýet ministrligi bilen ylalaşylan we '
    'Türkmenistanyň Döwlet gullugynyň başlygynyň 2024-nji ýylyň 22-nji maýyndaky 48 '
    'belgili buýrugy bilen tassyklanan “Gümrük edaralary hyzmatlary üçin tölegleriň '
    's/anawy we olaryň möçberleri” bilen tanyşdym.',
    'Harytlary ýükläp ugratmak we olary gümrük taýdan resmileşdirmek boýunça jogapkär',
]


def build_customs() -> Path:
    """Customs-clearance request letter (ARZA, Turkmen) — with truck table + boilerplate."""
    doc = _letter_doc()
    _addressee(doc, ['Aşgabat  gümrükhanasynyň', '“AKÝOL” gümrük nokadynyň', 'müdirine'])
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run('ARZA').bold = True
    doc.add_paragraph()
    _para(doc, _CUSTOMS_BODY, justify=True)
    doc.add_paragraph()

    cols = ['T/b', 'Ulag serişdeleriniň Belgisi', 'Harydyň ady', 'Orun sany', 'Harydyň agramy']
    fields = ['1', '{{ plate }}', '{{ product }}', '{{ boxes }}', '{{ gross }} kg.']
    table = doc.add_table(rows=2, cols=len(cols))
    table.style = 'Table Grid'
    _col_widths(table, [Cm(1.3), Cm(5), Cm(2.5), Cm(2.2), Cm(4)])
    for i, name in enumerate(cols):
        _set_cell(table.rows[0].cells[i], name, bold=True, size=11)
    for i, val in enumerate(fields):
        _set_cell(table.rows[1].cells[i], val, size=11)
    doc.add_paragraph()

    for para in _CUSTOMS_PARAS:
        _para(doc, para, justify=True)
    doc.add_paragraph()

    sign = doc.add_table(rows=1, cols=2)
    _col_widths(sign, [Cm(6), Cm(11)])
    _set_cell(sign.cell(0, 0), 'Telekeçi', bold=True, size=12)
    _set_cell(sign.cell(0, 1), '{{ seller_name }}', bold=True, size=12)

    out = OUT_DIR / 'customs_tk.docx'
    doc.save(out)
    return out


def main() -> None:
    for lang in ('ru', 'en'):
        print(f'wrote {build(lang)}')
    for builder in (build_ct1, build_fito, build_customs):
        print(f'wrote {builder()}')


if __name__ == '__main__':
    main()
