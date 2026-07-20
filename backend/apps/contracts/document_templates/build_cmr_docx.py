"""Build the Word CMR overlay templates (``cmr_ru_docx.docx`` / ``cmr_en_docx.docx``).

The CMR is a print overlay onto the pre-printed official 24-box form. The primary
template is the geometry-preserving ``.xlsx`` (see ``build_cmr_xlsx.py``); this
builds a **Word** counterpart for users who want to edit the document before
printing.

LibreOffice cannot convert a spreadsheet to Word (it refuses xlsx→docx), so the
Word variant can't be a conversion — it has to reproduce the grid natively. To
avoid inventing a second layout that could drift, the geometry is **derived from
the xlsx template itself**: each Excel column width / row height is converted to
centimetres, multiplied by the sheet's print scale, and emitted as a borderless
Word table of the same shape. Field placement reuses the very same
``_CMR_OVERLAY_CELLS`` coordinate map the xlsx overlay uses, so both formats put
every value in the same box.

CAVEAT: Word tables approximate — they do not guarantee — Excel's exact print
registration. The ``.xlsx`` remains the reference for printing onto the blank
form; this Word variant must be verified with a physical test print before it is
trusted for that purpose.

Run once to (re)create the committed templates:

    python -m apps.contracts.document_templates.build_cmr_docx
"""
from pathlib import Path

import openpyxl
from docx import Document
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

from apps.contracts.services.document_context import _CMR_OVERLAY_CELLS

OUT_DIR = Path(__file__).resolve().parent

# The printed CMR grid: columns A–N, rows 1–54 (matches the xlsx print area).
LAST_COL = 14
LAST_ROW = 54

# Excel's default column width / row height, used when a dimension is unset.
DEFAULT_COL_WIDTH = 8.43
DEFAULT_ROW_HEIGHT = 15.75

SOURCE = {'ru': 'cmr_ru.xlsx', 'en': 'cmr_en.xlsx'}
OUT_NAME = {'ru': 'cmr_ru_docx.docx', 'en': 'cmr_en_docx.docx'}


def _col_cm(width: float, scale: float) -> float:
    """Excel column width (character units) → centimetres at the print scale."""
    pixels = width * 7 + 5
    return (pixels * 2.54 / 96) * scale


def _row_cm(height: float, scale: float) -> float:
    """Excel row height (points) → centimetres at the print scale."""
    return (height / 72 * 2.54) * scale


def _no_borders(table) -> None:
    """Strip every border from the table — it overlays a pre-printed form."""
    tbl_pr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'nil')
        borders.append(el)
    tbl_pr.append(borders)


def _fixed_layout(table) -> None:
    """Pin the table to fixed layout so the cell widths are honoured as written."""
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    tbl_pr.append(layout)


def build(lang: str) -> Path:
    sheet = openpyxl.load_workbook(OUT_DIR / SOURCE[lang])
    ws = sheet.active
    scale = (ws.page_setup.scale or 100) / 100

    widths = []
    for idx in range(1, LAST_COL + 1):
        letter = openpyxl.utils.get_column_letter(idx)
        dim = ws.column_dimensions.get(letter)
        widths.append(_col_cm(dim.width if dim and dim.width else DEFAULT_COL_WIDTH, scale))

    heights = []
    for row in range(1, LAST_ROW + 1):
        dim = ws.row_dimensions.get(row)
        heights.append(_row_cm(dim.height if dim and dim.height else DEFAULT_ROW_HEIGHT, scale))

    doc = Document()
    section = doc.sections[0]
    section.page_width, section.page_height = Cm(21), Cm(29.7)
    # Excel stores margins in inches; mirror them so the grid starts where the
    # spreadsheet's does.
    margins = ws.page_margins
    section.left_margin = Cm(margins.left * 2.54)
    section.right_margin = Cm(margins.right * 2.54)
    section.top_margin = Cm(margins.top * 2.54)
    section.bottom_margin = Cm(margins.bottom * 2.54)
    doc.styles['Normal'].font.size = Pt(8)

    table = doc.add_table(rows=LAST_ROW, cols=LAST_COL)
    _no_borders(table)
    _fixed_layout(table)

    for r, row in enumerate(table.rows):
        row.height = Cm(heights[r])
        row.height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        for c, cell in enumerate(row.cells):
            cell.width = Cm(widths[c])
            para = cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)

    # Group the fields by row so each one can claim the free columns to its right.
    by_row: dict[int, list[tuple[int, str]]] = {}
    for coord, field in _CMR_OVERLAY_CELLS[lang].items():
        col_letter, row_idx = openpyxl.utils.cell.coordinate_from_string(coord)
        col_idx = openpyxl.utils.column_index_from_string(col_letter)
        if row_idx > LAST_ROW or col_idx > LAST_COL:
            raise ValueError(f'{lang}: {coord} falls outside the A1:N{LAST_ROW} grid')
        by_row.setdefault(row_idx, []).append((col_idx, field))

    # Place each field's Jinja tag in the SAME cell the xlsx overlay writes to.
    # In Excel a long value simply overflows into the blank cells to its right; a
    # Word table cell would clip it instead, so merge each field across the columns
    # up to the next field in that row (or the grid edge) to reproduce the spill.
    for row_idx, entries in by_row.items():
        entries.sort()
        row = table.rows[row_idx - 1]
        for i, (col_idx, field) in enumerate(entries):
            end_col = entries[i + 1][0] - 1 if i + 1 < len(entries) else LAST_COL
            if end_col > col_idx:
                row.cells[col_idx - 1].merge(row.cells[end_col - 1])
            cell = row.cells[col_idx - 1]
            cell.text = ''
            cell.paragraphs[0].add_run(f'{{{{ {field} }}}}').font.size = Pt(8)

    out = OUT_DIR / OUT_NAME[lang]
    doc.save(out)
    return out


def main() -> None:
    for lang in ('ru', 'en'):
        print(f'wrote {build(lang)}')


if __name__ == '__main__':
    main()
